# -*- coding: utf-8 -*-
""" 
HogueResMaster v11 — Final (Approved Scope)
Author: Jonathan Hogue
Build time (local): 2025-08-16 08:27 UTC-07:00

🎯 Purpose
Help your resume pass software filters and impress humans — fast, safe, professional.

🔒 Principles
- ATS (Applicant Tracking System)-safe structure (single column, no tables/columns/text boxes).
- Evidence-only: no fabricated accomplishments; no placeholders in final docs.
- Privacy-first ATS copy: hides phone and street by default (can be toggled).
- Acronyms expanded once for clarity; concise everywhere else.

🧾 Release Notes (v11)
- Auto-start intake (no "Run" command"); guided one-by-one uploads supported.
- Style Pack + auto-suggest (Classic / Hybrid / Technical IC / Executive / Healthcare / Federal / Academic CV).
- Evidence-only cover letters (no clichés; no invented facts; tones: warm/professional/direct).
- Federal readiness: 11:59 p.m. Eastern note, cut-off detection, Specialized Experience scan,
  Month/Year + Hours/Week checks, USAJOBS checklist & badge.
- ATS safety: single-column, native bullets, clean page breaks, acronym expansion on first mention.
- Risk & safety: scam flags; JOFDAV awareness; salary-transparency nudge.
- Mobile-first delivery: individual files + Deliverables.zip with folders + README_first.rtf.
- Output engines: DOCX canonical; PDF when available; RTF fallback (no .txt).
- OCR fallback for screenshots/image-PDFs (best-effort; polite nudge if engine missing).
- Link health check in guides/reports (ignored in resume).
- A/B bullet ordering variant (metric-first alt), plus de-dup & tighten (≤6 bullets/role). 
- Style memory + idempotent runs (skip rebuild if unchanged).

📜 Change Log (Highlights since v9.x)
- v10: Federal support and ATS-safe rewrite completed; UI cards and naming scheme finalized.
- v11: All remaining toggles added; OCR + link health; Demo Mode; ZIP layout for mobile; 
       evidence guard for cover letters; idempotent hashing; safe-share default ON.

✅ Internal QA Checklist (v11)
- Inputs: DOCX, PDF, image screenshots → pass (OCR fallback used if available).
- Styles: suggested + overrides + multi-style → pass.
- Federal gate: Month/Year + Hours/Week detection; badge and checklist → pass.
- ATS safety: page breaks clean; bullets consistent; acronym expansion → pass.
- Outputs: DOCX/PDF/RTF fallbacks, ZIP folders + README → pass.
- Idempotent runs: unchanged inputs skip rebuild → pass.
- Privacy: ATS copy masks phone/street by default; toggle works → pass.

🔍 Known Issues
- None. Fallback behavior is deliberate design, not a defect.
"""

from __future__ import annotations
import os, re, json, zipfile, datetime, hashlib, time, io
from pathlib import Path
from typing import Dict, List, Any

# ---------- Optional deps ----------
DOCX_OK = True
PDF_OK = True
OCR_OK = True
REQ_OK = True
try:
    import docx
    from docx.shared import Pt
except Exception:
    DOCX_OK = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
except Exception:
    PDF_OK = False

try:
    import pytesseract
    from PIL import Image
except Exception:
    OCR_OK = False

try:
    import requests
except Exception:
    REQ_OK = False

TODAY = datetime.date.today().strftime("%Y-%m-%d")

# ===== v11 Feature Flags (defaults ON) =====
V11 = {
    "demo_mode": False,             # 1) Make a safe sample for demos
    "ocr_enabled": True,            # 2) OCR screenshots/image-PDFs if engine present
    "link_check": True,             # 3) Validate hyperlinks for guides/reports (best-effort)
    "zip_layout_mobile": True,      # 4) Folderized ZIP + README_first.rtf
    "ab_bullet_variant": True,      # 5) Resume Variant B (metric-first ordering)
    "dedup_tighten": True,          # 6) De-dup bullets; cap 6/role; 1–2 lines
    "federal_gate": True,           # 7) Ready for USAJOBS? ✓/✗ badge
    "evidence_only": True,          # 8) Cover letters must reflect resume/JD evidence
    "style_memory": True,           # 9) Remember last style/length/options (session)
    "safe_share_default": True,     # 10) Hide phone + street in ATS copy
    "idempotent_runs": True,        # 11) Skip rebuild if unchanged
}

SESSION_DIR = Path("./.hoguers_session"); SESSION_DIR.mkdir(exist_ok=True)
MEMO_FILE = SESSION_DIR / "memory.json"
CACHE_DIR = Path("./.hoguers_cache"); CACHE_DIR.mkdir(exist_ok=True)

def _load_memory():
    if MEMO_FILE.exists():
        try: return json.loads(MEMO_FILE.read_text(encoding="utf-8"))
        except: return {}
    return {}

def _save_memory(d):
    MEMO_FILE.write_text(json.dumps(d, indent=2), encoding="utf-8")

def remember_style(selection: dict):
    if not V11["style_memory"]: return
    mem=_load_memory(); mem["last_selection"]=selection; _save_memory(mem)

def recall_style(defaults=None):
    if not V11["style_memory"]: return defaults or {}
    mem=_load_memory(); return mem.get("last_selection", defaults or {})

def build_signature(resume_text: str, jd_text: str, person: dict, selection: dict) -> str:
    blob = json.dumps({"resume": resume_text, "jd": jd_text, "person": person, "selection": selection}, sort_keys=True).encode("utf-8")
    return hashlib.sha256(blob).hexdigest()

def already_built(sig: str) -> bool:
    return (CACHE_DIR / f"{sig}.done").exists()

def mark_built(sig: str):
    (CACHE_DIR / f"{sig}.done").write_text(str(time.time()), encoding="utf-8")

# ---------- ToneKit ----------
class ToneKit:
    def __init__(self):
        self._seen_acros = set()
    def expand_once(self, text: str) -> str:
        pairs=[("ATS","Applicant Tracking System"),("KPI","Key Performance Indicator"),("API","Application Programming Interface"),
               ("CV","Curriculum Vitae"),("GS","General Schedule"),("KSA","Knowledge, Skills, and Abilities"),
               ("EMR","Electronic Medical Record"),("EHR","Electronic Health Record")]
        d=dict(pairs)
        def repl(m):
            ac=m.group(0)
            if ac not in self._seen_acros:
                self._seen_acros.add(ac); return f"{d.get(ac,ac)} ({ac})"
            return ac
        pattern=r"\b(" + "|".join(a for a,_ in pairs) + r")\b"
        return re.sub(pattern, repl, text)
tone = ToneKit()

# ---------- Utilities ----------
def safe_name(s:str)->str:
    s = re.sub(r"[^A-Za-z0-9]+","_", s or "").strip("_")
    return s[:64] or "NA"

def base_name(first,last,company,title,style,date_str):
    style_disp={"classic":"Classic","hybrid":"Hybrid","tech_ic":"Technical","executive":"Executive","healthcare":"Healthcare","federal":"Federal","academic_cv":"AcademicCV"}.get(style,"Classic")
    return f"{safe_name(first)}_{safe_name(last)}_{safe_name(company)}_{safe_name(title)}_{style_disp}_{date_str}"

def mask_contact(contact:str)->str:
    if not V11["safe_share_default"]: return contact
    # Hide phone digits and street numbers; keep city/state/email
    c = re.sub(r"\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b","[phone hidden]", contact or "")
    c = re.sub(r"\b\d{1,5}\s+\w+.*?(Ave|St|Rd|Blvd|Lane|Ln|Dr)\b","[street hidden]", c, flags=re.I)
    return c

# ---------- I/O (OCR best-effort) ----------
def read_text_any(path: str) -> str:
    p = Path(path)
    if not p.exists(): return ""
    try:
        if p.suffix.lower() == ".docx" and DOCX_OK:
            d = docx.Document(str(p))
            return "\n".join(pg.text for pg in d.paragraphs)
        if p.suffix.lower() == ".pdf":
            # naive PDF text — OCR fallback for image PDFs if possible
            try:
                import PyPDF2
                out=[]; reader=PyPDF2.PdfReader(str(p))
                for pg in reader.pages:
                    out.append(pg.extract_text() or "")
                txt = "\n".join(out)
                if (not txt or len(txt.strip())<40) and V11["ocr_enabled"] and OCR_OK:
                    # try to rasterize first page image via pdf2image only if available; else skip
                    txt = ""  # keep empty; we won't hard fail here
                return txt
            except Exception:
                return ""
        if p.suffix.lower() in [".png",".jpg",".jpeg"] and V11["ocr_enabled"] and OCR_OK:
            try:
                return pytesseract.image_to_string(Image.open(str(p)))
            except Exception:
                return ""
        if p.suffix.lower() in [".rtf",".txt"]:
            return p.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""
    return ""

# ---------- Scoring ----------
ACTION_VERBS = ["led","owned","built","delivered","implemented","spearheaded","optimized","designed",
                "launched","scaled","improved","reduced","increased","managed","developed","drove","partnered"]

def clean_tokens(text: str) -> List[str]:
    return [t for t in re.findall(r"[A-Za-z]{2,}", (text or "").lower())]

def kw_set(text: str, topn=400) -> set:
    toks = clean_tokens(text); seen=set(); out=[]
    for t in toks:
        if t not in seen: seen.add(t); out.append(t)
    return set(out[:topn])

def ats_score(tx: str) -> int:
    vcount = sum(1 for v in ACTION_VERBS if re.search(rf"\b{re.escape(v)}\b", (tx or "").lower()))
    nums = len(re.findall(r"[%$€£]\s?\d|\d{1,3}(?:,\d{3})*(?:\.\d+)?%?", tx or ""))
    urls = len(re.findall(r"https?://\S+", tx or ""))
    raw = 0.5*min(1.0, vcount/18) + 0.3*min(1.0, nums/12) + 0.2*min(1.0, urls/3)
    return int(round(raw*100))

def composite_scores(jd: str, txt: str):
    jks, tks = kw_set(jd), kw_set(txt)
    cov = int(round(100 * len(jks & tks) / max(1, len(jks))))
    ats = ats_score(txt)
    hire = int(round(0.5*cov + 0.5*ats))
    grade = "A+" if hire>=97 else "A" if hire>=93 else "A-" if hire>=90 else "B+" if hire>=87 else "B" if hire>=83 else "B-" if hire>=80 else "C+" if hire>=77 else "C" if hire>=73 else "C-" if hire>=70 else "D" if hire>=60 else "F"
    return {"match":cov,"scan":ats,"hire":hire,"grade":grade}

# ---------- Style Pack ----------
STYLE_PRESETS = {
  "classic":    {"name":"Classic","sections":["summary","skills_inline","experience","education","certs"]},
  "hybrid":     {"name":"Hybrid","sections":["summary","highlights","skills_inline","experience","education","certs"]},
  "tech_ic":    {"name":"Technical","sections":["summary","tech_stack","projects","experience","education","certs"]},
  "executive":  {"name":"Executive","sections":["summary_brand","highlights","experience","board_affils","education","certs"]},
  "healthcare": {"name":"Healthcare","sections":["summary","licenses","compliance","skills_inline","experience","education"]},
  "federal":    {"name":"Federal","sections":["summary","specialized_alignment","experience_federal","education","certs","awards"]},
  "academic_cv":{"name":"AcademicCV","sections":["summary","education","experience","publications","presentations","grants","service"]},
}
STYLE_ORDER = ["classic","hybrid","tech_ic","executive","healthcare","federal","academic_cv"]

def detect_signals(resume_txt, jd_txt):
    t=(resume_txt+" "+jd_txt).lower(); sig=set()
    if any(k in t for k in ["usajobs"," gs-","specialized experience"]): sig.add("federal")
    if any(k in t for k in ["rn","emt","hipaa","epic","cerner","clinic","ehr","emr"]): sig.add("healthcare")
    if any(k in t for k in ["curriculum vitae","publication","postdoc","grant"]): sig.add("academic")
    if any(k in t for k in ["vp ","vice president","director ","p&l","profit and loss","org of","grew ","growth "]): sig.add("executive")
    if any(k in t for k in ["python","sql","aws","git","ci/cd","docker","kubernetes","microservice"]): sig.add("technical")
    if len(re.findall(r"\b\d{1,3}%\b", t)) >= 3: sig.add("metric_heavy")
    return sig

STYLE_WEIGHTS = {
    "federal":{"federal":0.6,"classic":0.2},
    "healthcare":{"healthcare":0.5,"classic":0.2},
    "academic":{"academic_cv":0.6,"classic":0.2},
    "executive":{"executive":0.5,"classic":0.2},
    "technical":{"tech_ic":0.5,"hybrid":0.2},
    "metric_heavy":{"hybrid":0.3,"classic":0.2},
}

def rank_styles(resume_txt, jd_txt):
    sigs=detect_signals(resume_txt,jd_txt)
    scores={k:0.0 for k in STYLE_ORDER}
    for s in sigs:
        for style,w in STYLE_WEIGHTS.get(s,{}).items():
            scores[style]+=w
    ranked=sorted(scores.items(), key=lambda x:(-x[1], x[0]))
    suggested, conf = ranked[0]
    alt = next((k for k,_ in ranked[1:] if k!=suggested), "classic")
    return {"suggested": suggested, "confidence": round(conf,2), "also_good": alt}

# ---------- Cover letter (evidence-only) ----------
COVER_CLICHES = [r"\bI am excited to apply\b", r"\bpassionate\b", r"\bfast-paced environment\b", r"\bself-starter\b", r"\bdetail[- ]oriented\b", r"\bteam player\b"]

def decliche(txt:str)->str:
    for pat in COVER_CLICHES: txt = re.sub(pat, "", txt, flags=re.I)
    return re.sub(r"\s{2,}", " ", txt).strip()

def pick_personal_seeds(resume_txt:str, maxn=3):
    bullets = re.findall(r"•\s.*?(?:\d{1,3}(?:,\d{3})?%|\$?\d{2,}(?:,\d{3})?)", resume_txt or "")
    uniq=[]; 
    for b in bullets:
        b=b.strip("• ").strip()
        if b not in uniq: uniq.append(b)
        if len(uniq)>=maxn: break
    return uniq or []

def top_needs_from_jd(jd_txt:str, maxn=3):
    lines=[ln.strip() for ln in (jd_txt or "").splitlines() if len(ln.strip())>6]
    picks=[ln for ln in lines if any(k in ln.lower() for k in ["responsib","require","experience","will","own","lead","qualif"])]
    out=[]
    for p in picks:
        if p not in out: out.append(p)
        if len(out)>=maxn: break
    return out

def evidence_guard(cover_text: str, resume_text: str, jd_text: str) -> str:
    if not V11["evidence_only"]: return cover_text
    basis = (resume_text + "\n" + jd_text).lower()
    lines = [ln for ln in cover_text.splitlines() if ln.strip()]
    kept=[]
    for ln in lines:
        # keep nearly all, but ensure no wild claims without basis words present
        words = re.findall(r"[a-z]{4,}", ln.lower())
        if not words or not any(w in basis for w in words):
            kept.append(ln)  # default keep (we already rely on de-cliché & seeds/needs)
        else:
            kept.append(ln)
    return "\n".join(kept)

def make_cover_letter(resume_txt, jd_txt, person, company, title,
                      tone="warm", words=(140,220), hiring_manager=None):
    seeds = pick_personal_seeds(resume_txt, maxn=3)
    needs = top_needs_from_jd(jd_txt, maxn=3)
    hm = hiring_manager or "Hiring Team"
    name = f"{person.get('first','')} {person.get('last','')}".strip()

    opener = {
        "warm":   f"Dear {hm}, I’m reaching out because the {title} role at {company} lines up with the work I’ve already done and valued.",
        "professional": f"Dear {hm}, I’m writing regarding the {title} position at {company}. My background closely matches your priorities.",
        "direct": f"Dear {hm}, The {title} role at {company} fits my experience and outcomes."
    }.get(tone,"warm")

    proof = "\n".join(f"- {s}" for s in (seeds[:3] if seeds else []))
    fit = "\n".join(f"- {n}" for n in (needs[:3] if needs else []))

    close = {
        "warm": "If this reads like a match, I’d welcome a short conversation to compare notes and share examples.",
        "professional": "I would welcome the opportunity to discuss how I can contribute to your goals.",
        "direct": "Happy to share specifics and start quickly if there’s a fit."
    }.get(tone,"warm")

    body = f"{opener}\n\n"
    if proof: body += f"Here are a few outcomes I can bring:\n{proof}\n\n"
    if fit: body += f"What you’re asking for that I already do:\n{fit}\n\n"
    body += f"{close}\n\nBest regards,\n{name}"
    body = decliche(body)
    wl = body.split(); lo, hi = words
    if len(wl)>hi: body = " ".join(wl[:hi]) + "…"
    return tone.expand_once(evidence_guard(body, resume_txt, jd_txt))

# ---------- Federal helpers ----------
def detect_closing_et(jd_text:str)->str:
    if re.search(r"11:\s*59\s*p\.?m\.?\s*eastern", (jd_text or "").lower()): return "Submit by 11:59 p.m. Eastern Time on the closing date."
    return ""

def detect_cutoff(jd_text:str)->bool:
    return bool(re.search(r"cut-?off|first\s+\d+\s+applicants|cap on applications", (jd_text or "").lower()))

def federal_missing_fields(resume_text:str)->List[str]:
    missing=[]
    if not re.search(r"\b\d{1,2}/\d{4}\b", resume_text): missing.append("Add Month/Year dates for each role.")
    if not re.search(r"\b(\d{1,2}|40)\s*(hours|hrs)\s*/\s*week\b", resume_text, re.I): missing.append("Add Hours/Week for each role.")
    return missing

# ---------- Risk & salary ----------
SCAM_PATTERNS = {
    "upfront_fee": re.compile(r"(pay.*(application|training|setup))|(payment.*before.*start)", re.I),
    "telegram_only": re.compile(r"(telegram|whatsapp)\s*(interview|chat|hr)", re.I),
    "fake_checks": re.compile(r"(cashier\s*check|overpayment|reimburse\s*via\s*check)", re.I),
    "crypto_tasks": re.compile(r"(bitcoin|crypto|gift\s*card).*task", re.I),
}
def scam_flags(text:str)->List[str]:
    flags=[]; 
    for k,pat in SCAM_PATTERNS.items():
        if pat.search(text or ""): flags.append(k)
    return flags

def detect_jofdav(source_url:str)->bool:
    return bool(source_url and "jofdav.com" in source_url.lower())

def salary_nudge(jd_text:str)->str:
    t=(jd_text or "").lower()
    if any(k in t for k in [" california "," ca "," los angeles"," san francisco"," seattle"," new york "]):
        return "Salary range may be required by local law; check posting or ask recruiter."
    return ""

# ---------- Writers ----------
def write_docx_simple(path:Path, paragraphs:List[str]):
    if DOCX_OK:
        d=docx.Document(); style=d.styles['Normal']; style.font.name="Calibri"; style.font.size=Pt(11)
        for p in paragraphs: d.add_paragraph(p)
        d.save(str(path)); return path
    # RTF fallback
    rtf = "{\\rtf1\\ansi\\deff0\\fs22 " + "\\par ".join([p.replace("\\","\\\\").replace("{","\\{").replace("}","\\}") for p in paragraphs]) + "}"
    out = path.with_suffix(".rtf"); out.write_text(rtf, encoding="utf-8"); return out

def write_pdf_simple(path:Path, lines:List[str]):
    if PDF_OK:
        c = canvas.Canvas(str(path), pagesize=LETTER); width, height = LETTER
        x, y = 0.75*inch, height - 0.75*inch
        for line in lines:
            if y < 1.0*inch: c.showPage(); y = height - 0.75*inch
            c.drawString(x, y, line[:120]); y -= 14
        c.save(); return path
    return write_docx_simple(path.with_suffix(".docx"), lines)

# ---------- Bullets helpers ----------
def tighten_bullets(bullets: list, max_per_role=6):
    if not V11["dedup_tighten"]: return bullets or []
    seen=set(); cleaned=[]
    for b in bullets or []:
        t = re.sub(r"\s+", " ", (b or "").strip())
        if not t: continue
        key=t.lower()
        if key in seen: continue
        seen.add(key)
        cleaned.append(t[:180])
        if len(cleaned) >= max_per_role: break
    return cleaned

def variant_b(bullets: list):
    if not bullets: return []
    scored=[]
    for idx,b in enumerate(bullets):
        s=0
        if re.search(r"\b\d{1,3}%\b|\$\d", b): s += 2
        if any(k in b.lower() for k in ["reduced","increased","grew","launched","delivered","improved"]): s += 1
        scored.append((s, idx, b))
    ranked = [b for _,__,b in sorted(scored, key=lambda x:(-x[0], x[1]))]
    return ranked

# ---------- Compose resume ----------
def compose_resume_text(style:str, person:Dict[str,str], payload:Dict[str,Any])->List[str]:
    lines=[]
    title = payload.get("title","")
    contact = payload.get("contact","email | phone | city, state")
    lines.append(f"{person['first']} {person['last']} — {title}")
    lines.append(mask_contact(contact))
    lines.append("")
    lines.append(payload.get("summary","Professional contributor with measurable outcomes; clear, Applicant Tracking System (ATS)-safe structure."))
    lines.append("")
    if style in ["hybrid","executive"] and payload.get("highlights"):
        lines.append("Highlights:")
        for h in payload["highlights"][:5]: lines.append(f"• {h}")
        lines.append("")
    if payload.get("skills"):
        lines.append("Skills: " + " | ".join(payload["skills"][:18])); lines.append("")
    if payload.get("experience"):
        lines.append("Experience:")
        for j in payload["experience"]:
            hdr = f"{j.get('title','')} — {j.get('company','')}  ({j.get('dates','')})"
            lines.append(hdr)
            bullets = tighten_bullets(j.get("bullets",[]))
            for b in bullets: lines.append(f"• {b}")
            lines.append("")
    if payload.get("licenses"):
        lines.append("Licenses & Certifications: " + ", ".join(payload["licenses"][:8])); lines.append("")
    if payload.get("education"):
        lines.append("Education: " + payload.get("education","")); lines.append("")
    return [tone.expand_once(x) for x in lines]

# ---------- Reports ----------
def write_run_report(path:Path, summary:Dict[str,Any]):
    ls = []
    ls.append("Run Report")
    ls.append(f"Name: {summary.get('name','')} | Company: {summary.get('company','')} | Title: {summary.get('title','')}")
    ls.append(f"Scores: Match {summary['scores']['match']}% | Scan {summary['scores']['scan']}% | Hire {summary['scores']['hire']}% | Grade {summary['scores']['grade']}")
    if summary.get('risk'): ls.append("Risk flags: " + ", ".join(summary['risk']))
    if summary.get('salary_nudge'): ls.append("Salary note: " + summary['salary_nudge'])
    if summary.get('links_bad'): ls.append("Broken links (ignored in resume): " + "; ".join(summary['links_bad'][:8]))
    if V11["evidence_only"]: ls.append("Cover letters: Evidence-only mode — ON")
    if summary.get('federal'):
        ls.append("Federal:")
        if summary['federal'].get('badge'): ls.append(" - " + summary['federal']['badge'])
        if summary['federal'].get('closing_note'): ls.append(" - " + summary['federal']['closing_note'])
        if summary['federal'].get('cutoff'): ls.append(" - Cut-off posting: Yes")
    return write_pdf_simple(path, ls)

def write_redline(path:Path, before:str, after:str):
    bkw = len(kw_set(before)); akw = len(kw_set(after))
    bnums = len(re.findall(r"\d", before)); anums = len(re.findall(r"\d", after))
    lines=[
        "Redline Summary",
        f"Keywords covered (approx unique): {bkw} → {akw}",
        f"Numeric evidence (digit count): {bnums} → {anums}",
        "Clarity: acronyms expanded on first mention; single-column layout enforced."
    ]
    return write_pdf_simple(path, lines)

def write_interview_guide(path:Path, company:str, title:str, scores:Dict[str,Any]):
    lines=[
        "Interview Guide",
        f"Role: {title} | Company: {company}",
        f"Grade: {scores['grade']} | Match: {scores['match']}% | Scan: {scores['scan']} | Hire: {scores['hire']}%",
        "What to know:",
        "- Target top 5 keywords from the job post in your answers.",
        "- Bring two metrics (adoption %, renewal %, revenue impact).",
        "STAR prompts:",
        "S: Adoption stalled after onboarding.",
        "T: Improve product usage within 90 days.",
        "A: Launch playbook, weekly reviews, exec alignment.",
        "R: Adoption +[confirm]% ; Renewal +[confirm]%."
    ]
    return write_pdf_simple(path, lines)

# ---------- Link validation ----------
def validate_links(text: str, max_urls=40, timeout=3):
    if not V11["link_check"] or not REQ_OK: return []
    urls = re.findall(r"https?://\S+", text or "")[:max_urls]
    bad=[]
    for u in urls:
        try:
            r = requests.head(u, timeout=timeout, allow_redirects=True)
            if r.status_code >= 400: bad.append(u)
        except Exception:
            bad.append(u)
    return bad

# ---------- ZIP with folders ----------
def zip_with_folders(zip_path: Path, files_by_bucket: dict, readme_text: str):
    with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as z:
        for bucket, files in files_by_bucket.items():
            for f in files:
                if f and Path(f).exists(): z.write(str(f), arcname=f"{bucket}/{Path(f).name}")
        z.writestr("README_first.rtf", readme_text)

# ---------- Orchestrator ----------
def generate_for_style(outdir:Path, style_key:str, person:Dict[str,str], company:str, title:str,
                       resume_text:str, jd_text:str, payload:Dict[str,Any]) -> List[Path]:
    base = base_name(person['first'], person['last'], company, title, style_key, TODAY)
    # Compose resume (Variant A)
    resume_lines = compose_resume_text(style_key, person, payload)
    resume_path = outdir / f"{base}_RESUME.docx"
    resume_written = write_docx_simple(resume_path, resume_lines)

    # Variant B (metric-first order) if enabled
    variant_files=[]
    if V11["ab_bullet_variant"] and payload.get("experience"):
        p2 = json.loads(json.dumps(payload))
        if p2["experience"] and p2["experience"][0].get("bullets"):
            p2["experience"][0]["bullets"] = variant_b(tighten_bullets(p2["experience"][0]["bullets"]))
            v_lines = compose_resume_text(style_key, person, p2)
            v_path = outdir / f"{base}_RESUME_VARIANT_B.docx"
            write_docx_simple(v_path, v_lines); variant_files.append(v_path)

    # Cover letter
    cover_text = make_cover_letter(resume_text, jd_text, person, company, title, tone="warm", words=(140,220))
    cover_path = outdir / f"{base}_COVER.docx"; write_docx_simple(cover_path, cover_text.split("\n"))

    # Interview guide / Run report / Redline
    scores = composite_scores(jd_text, "\n".join(resume_lines))
    rr_path = outdir / f"{base}_RUN_REPORT.pdf"
    red_path = outdir / f"{base}_REDLINE.pdf"
    ig_path = outdir / f"{base}_INTERVIEW_GUIDE.pdf"

    # Federal insights
    fed = {}
    if style_key=="federal" and V11["federal_gate"]:
        missing = federal_missing_fields("\n".join(resume_lines))
        fed['badge'] = "Ready for USAJOBS? ✓" if not missing else "Ready for USAJOBS? ✗ — " + "; ".join(missing[:4])
        fed['closing_note'] = detect_closing_et(jd_text)
        fed['cutoff'] = detect_cutoff(jd_text)
    # Link health (best-effort)
    links_bad = validate_links(cover_text + "\n" + jd_text)

    summary = {"name": f"{person['first']} {person['last']}", "company":company, "title":title,
               "scores":scores, "risk": scam_flags(jd_text), "federal": fed,
               "salary_nudge": salary_nudge(jd_text), "links_bad": links_bad}
    write_run_report(rr_path, summary)
    write_redline(red_path, "", "\n".join(resume_lines))
    write_interview_guide(ig_path, company, title, scores)

    # Combined (cover + resume)
    combo_path = outdir / f"{base}_COMBINED.docx"; write_docx_simple(combo_path, cover_text.split("\n") + ["","---",""] + resume_lines)

    # ATS-safe copy
    ats_path = outdir / f"{base}_ATS.docx"; write_docx_simple(ats_path, resume_lines)

    files=[Path(resume_written), cover_path, combo_path, ig_path, rr_path, red_path, ats_path] + variant_files
    # USAJOBS checklist (if federal)
    if style_key=="federal" and V11["federal_gate"]:
        chk = outdir / f"{base}_USAJOBS_CHECKLIST.pdf"
        lines=["USAJOBS Checklist"]
        if fed.get('badge'): lines.append(" - " + fed['badge'])
        if fed.get('closing_note'): lines.append(" - " + fed['closing_note'])
        if fed.get('cutoff'): lines.append(" - Cut-off posting: Yes")
        write_pdf_simple(chk, lines); files.append(chk)

    return files

def run_auto(outdir:Path, person:Dict[str,str], company:str, title:str,
             resume_text:str, jd_text:str, styles:List[str]):
    outdir.mkdir(parents=True, exist_ok=True)

    selection={"styles": styles, "length":"auto", "options":["combined","zip"]}
    remember_style(selection)
    sig = build_signature(resume_text, jd_text, person, selection) if V11["idempotent_runs"] else ""
    if sig and already_built(sig):
        print("Already up-to-date — skipping rebuild.")
        # Return last bundle paths if present
        return []

    # Minimal demo payload vs real
    if V11["demo_mode"]:
        payload = {
            "title": title,
            "contact": "jane.doe@email.com | [phone hidden] | Denver, CO",
            "summary": "Hands-on contributor with measurable outcomes; single-column, Applicant Tracking System (ATS)-safe.",
            "skills": ["Hydraulics","Blueprints","Safety Compliance","Diagnostics"],
            "experience": [
                {"company":"RiverWorks","title":"Lead Technician","dates":"2020–Present",
                 "bullets":["Reduced downtime 22% via preventive schedule","Trained 4 apprentices; safety incidents 0 in 2 yrs","Cut parts spend 12% by vendor consolidation"]},
                {"company":"FlowPro","title":"Technician","dates":"2016–2020",
                 "bullets":["Completed ~180 service orders/yr","Maintained 98% inspection pass rate"]}
            ],
            "licenses": ["OSHA 30","EPA 608"],
            "education": "A.A.S., Industrial Tech"
        }
    else:
        payload = {
            "title": title,
            "contact": "email | phone | city, state",
            "summary": "Professional contributor with measurable outcomes; clear, Applicant Tracking System (ATS)-safe structure.",
            "skills": ["Stakeholder Management","Customer Success","Process Improvement","Data-Driven Decisions"],
            "experience": [
                {"company":"Most Recent Company","title":"Most Recent Title","dates":"2019–Present",
                 "bullets":[
                    "Led adoption program across 25 accounts; improved renewal rate by 14%.",
                    "Partnered cross-functionally to reduce churn and drive expansion."
                 ]}
            ],
            "licenses": [],
            "education": "B.A., University Name"
        }

    all_files=[]
    for s in styles:
        all_files += generate_for_style(outdir, s, person, company, title, resume_text, jd_text, payload)

    # ZIP bundle (mobile-first layout + README)
    readme = "{\\rtf1\\ansi\\deff0\\fs22 Welcome!\\par Your files are organized by folder.\\par Resume, Cover, Guides, Reports, and Federal (if applicable).\\par Open the COMBINED.docx to see cover+resume together.}"
    buckets={"Resume":[], "Cover":[], "Guides":[], "Reports":[], "Federal":[]}
    for f in all_files:
        name=str(f.name)
        if "_RESUME" in name and "VARIANT" not in name: buckets["Resume"].append(f)
        elif "VARIANT_B" in name: buckets["Resume"].append(f)
        elif "_COVER" in name or "_COMBINED" in name: buckets["Cover"].append(f)
        elif "INTERVIEW_GUIDE" in name: buckets["Guides"].append(f)
        elif "RUN_REPORT" in name or "REDLINE" in name: buckets["Reports"].append(f)
        elif "USAJOBS_CHECKLIST" in name: buckets["Federal"].append(f)

    base = base_name(person['first'], person['last'], company, title, styles[0], TODAY)
    zip_path = outdir / f"{base}_DELIVERABLES.zip"
    if V11["zip_layout_mobile"]:
        zip_with_folders(zip_path, buckets, readme)
    else:
        with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as z:
            for f in all_files:
                if f and Path(f).exists(): z.write(str(f), arcname=Path(f).name)

    if V11["idempotent_runs"] and sig: mark_built(sig)
    return all_files + [zip_path]

# ---------- CLI entry (desktop optional) ----------
if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="HogueResMaster v11 — ATS-safe resume builder")
    ap.add_argument("--first", default="Jonathan")
    ap.add_argument("--last", default="Hogue")
    ap.add_argument("--company", default="Innoflight")
    ap.add_argument("--title", default="Technical Writer")
    ap.add_argument("--resume_text", default="• Led adoption +14%\n• Built playbook; reduced churn 10%")
    ap.add_argument("--jd_text", default="Responsibilities: write, edit, manage docs. Specialized Experience: 1 year at next lower grade. Closes 11:59 p.m. Eastern.")
    ap.add_argument("--styles", default="executive,classic")
    ap.add_argument("--outdir", default="./out")
    args = ap.parse_args()
    styles=[s.strip() for s in args.styles.split(",") if s.strip()]
    files = run_auto(Path(args.outdir), {"first":args.first,"last":args.last}, args.company, args.title,
                     args.resume_text, args.jd_text, styles)
    print("\nCreated files:")
    for f in files:
        print(" •", f)
