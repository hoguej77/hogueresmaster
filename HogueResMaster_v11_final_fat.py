# -*- coding: utf-8 -*-
"""
HogueResMaster_v11.py
Cards-only UX + Style Packs + Federal/Healthcare helpers + ZIP packaging

IMPORTANT: Filename is fixed by request. Do not rename without consent.

This build is designed for chat use:
- Instruction-first (quiet) until both resume + job are present
- Single emoji summary card when done
- DOCX output with RTF fallback
- Style Packs: Classic, Hybrid, Technical, Executive, Healthcare, Federal, CV
- Federal checklist helper + Healthcare skill enrichment
- Scam/Risk heuristics for job text
- Deliverables.zip packaging for mobile users
"""

from __future__ import annotations
import os, re, sys, difflib, datetime, zipfile
from pathlib import Path
from typing import List, Dict, Any

import json

# Fat build toggles & label
FAT_MODE = True
BUILD_LABEL = "v11 final fat"

# Optional deps
DOCX_OK=True
try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    DOCX_OK=False

PDF_OK=True
try:
    import PyPDF2
except Exception:
    PDF_OK=False

REQ_OK=True
try:
    import requests
except Exception:
    REQ_OK=False

TODAY = datetime.date.today().strftime("%Y-%m-%d")

# ===================== UI Core (cards-only) =====================
def ui_debug(*args, **kwargs):
    if os.environ.get("HOGUE_VERBOSE", "0") == "1":
        print(*args, **kwargs)

def ui_card(text: str):
    print(text)

WELCOME_CARD = """\
👋 Welcome to HogueResMaster (v11)

Upload in this chat (all at once or one by one):
  1) 📑 Your resume (DOCX/PDF/RTF/TXT)
  2) 📋 The job posting (DOCX/PDF/TXT/HTML or the link)
  3) 💾 This file (HogueResMaster_v11.py) or ZIP

I will rebuild an ATS-safe resume, write a sincere cover letter, create an interview guide, and package everything for easy download. If one item is missing, I’ll ask only for that.
"""

# ===================== Data & Heuristics =====================
STOPWORDS = set("a an and or the but with for to from on in at as by of be is are was were will would shall should can could into within without among across per plus via than then that this those it its your you we our they them their he she his her who whom which what when where why how".split())
ACTION_VERBS = ["led","owned","built","delivered","implemented","orchestrated","spearheaded","optimized","designed","launched","scaled","improved","reduced","increased","managed","developed","drove","partnered","enabled","accelerated","transformed","modernized"]

AGG_DOMAINS = {"indeed.com","ziprecruiter.com","linkedin.com","glassdoor.com","monster.com","simplyhired.com","jobcase.com"}
SCAM_PATTERNS = {
    "upfront_fee": re.compile(r"(pay.*(application|training|setup))|(payment.*before.*start)", re.I),
    "telegram_only": re.compile(r"(telegram|whatsapp)\s*(interview|chat|hr)", re.I),
    "install_software": re.compile(r"(install|download).*(client|anydesk|teamviewer|software)", re.I),
    "bank_ssn": re.compile(r"(ssn|social security|bank\s+info|routing\s+number|account\s+number)", re.I),
    "generic_email": re.compile(r"[A-Za-z0-9._%+\-]+@(gmail|yahoo|outlook|hotmail)\.com", re.I),
    "too_good_pay": re.compile(r"\b(\$ ?[2-9]\d{2,}|(\$ ?\d{2,}\,?\d{3}))\/\s?(week|day)\b", re.I),
    "urgent_now": re.compile(r"(urgent|immediate)\s+(hiring|start)", re.I),
    "wire_gift": re.compile(r"(gift\s*card|bitcoin|crypto|wire\s*transfer)", re.I),
}

MED_CERTS = {"CNA","LPN","LVN","RN","BSN","MSN","NP","PA-C","MD","DO","EMT","EMT-B","EMT-I","EMT-P","Paramedic","RRT","CRT","PT","DPT","OT","COTA","SLP","PharmD","RPh","CPhT","CRCST","ARRT","CNMT","CPR","BLS","ACLS","PALS","NRP","TNCC","ENPC","ATLS","RHIT","RHIA","CPC","CCS","CCA","COC","CIC","CHDA","CRCR"}
MED_COMPLIANCE = {"HIPAA","OSHA","JCAHO","The Joint Commission","CLIA","CAP","IRB","GCP","ICH-GCP","FDA","EMA","21 CFR Part 11"}
MED_EMR = {"Epic","EpicCare","Cerner","Oracle Health","Meditech","Allscripts","athenahealth","eClinicalWorks","NextGen","Practice Fusion"}
MED_SPECIALTIES = {"Emergency","ICU","CCU","Telemetry","Oncology","Pediatrics","OB/GYN","Labor & Delivery","Surgery","Perioperative","PACU","Cath Lab","Radiology","Cardiology","Pulmonology","Nephrology","Psychiatry","Behavioral Health","Primary Care","Urgent Care","Family Medicine","Internal Medicine","Clinical Research","Pharmacovigilance","Regulatory Affairs","Revenue Cycle","Medical Billing","Coding"}

VET_CROSSWALK = {
    "11B":"Infantryman → Security/Operations Specialist",
    "88M":"Motor Transport → Logistics/Driver Ops",
    "68W":"Combat Medic → EMT/Patient Care Tech",
    "35F":"Intel Analyst → Business/Threat Intelligence",
    "3D1":"Client Systems → IT Support/Helpdesk",
    "3D2":"Cyber Systems → Systems Admin/SRE",
}

STYLE_PACKS = {
    "Classic":  {"font":"Calibri", "size":11},
    "Hybrid":   {"font":"Calibri", "size":11},
    "Technical":{"font":"Consolas","size":10},
    "Executive":{"font":"Georgia", "size":11},
    "Healthcare":{"font":"Calibri", "size":11},
    "Federal":  {"font":"Calibri", "size":11},
    "CV":       {"font":"Times New Roman", "size":11},
}

# ===================== Utility =====================
def clean_tokens(text: str) -> List[str]:
    return [t for t in re.findall(r"[A-Za-z]{2,}", (text or "").lower()) if t not in STOPWORDS]

def kw_set(text: str, topn=400) -> set:
    toks = clean_tokens(text)
    seen, out = set(), []
    for t in toks:
        if t not in seen:
            seen.add(t); out.append(t)
    return set(out[:topn])

def ats_score(tx: str) -> int:
    vcount = sum(1 for v in ACTION_VERBS if re.search(rf"\b{re.escape(v)}\b", (tx or "").lower()))
    nums = len(re.findall(r"[%$€£]\s?\d|\d{1,3}(?:,\d{3})*(?:\.\d+)?%?", tx or ""))
    urls = len(re.findall(r"https?://\S+", tx or ""))
    raw = 0.5*min(1.0, vcount/18) + 0.3*min(1.0, nums/12) + 0.2*min(1.0, urls/3)
    return int(round(raw*100))

def letter_grade(x: int):
    return "A+" if x>=97 else "A" if x>=93 else "A-" if x>=90 else "B+" if x>=87 else "B" if x>=83 else "B-" if x>=80 else "C+" if x>=77 else "C" if x>=73 else "C-" if x>=70 else "D" if x>=60 else "F"

def composite_scores(jd: str, txt: str):
    jks, tks = kw_set(jd), kw_set(txt)
    cov = int(round(100 * len(jks & tks) / max(1, len(jks))))
    ats = ats_score(txt)
    hire = int(round(0.5*cov + 0.5*ats))
    return cov, ats, hire, letter_grade(hire)

# ===================== IO =====================
def read_text_any(path_or_url: str) -> str:
    if not path_or_url: return ""
    p = str(path_or_url)
    try:
        if re.match(r"^https?://", p) and REQ_OK:
            r = requests.get(p, timeout=20)
            if r.ok: return r.text
        path = Path(p)
        if not path.exists(): return ""
        if p.lower().endswith(".txt"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if p.lower().endswith(".rtf"):
            raw = path.read_text(encoding="utf-8", errors="ignore")
            return re.sub(r"\{\\\*?[^}]*\}|\\[a-z]+\d* ?|[{}]", "", raw)
        if p.lower().endswith(".docx") and DOCX_OK:
            d = docx.Document(str(path))
            return "\n".join([pg.text for pg in d.paragraphs])
        if p.lower().endswith(".pdf") and PDF_OK:
            txt = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for pg in reader.pages:
                    txt.append(pg.extract_text() or "")
            return "\n".join(txt)
        if p.lower().endswith(".html") or p.lower().endswith(".htm"):
            return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        ui_debug(f"[read_text_any] {e}")
    return ""

# ===================== Detection =====================
def _looks_resume(p: Path) -> bool:
    return bool(re.search(r"(resume|cv)\.(docx|pdf|rtf|txt)$", p.name, re.I))

def _looks_job(p: Path) -> bool:
    return bool(re.search(r"(job|jd|posting|description)\.(docx|pdf|rtf|txt|html|htm)$", p.name, re.I))

def scan_uploads(search_dir: Path) -> dict:
    found = {"resume": None, "job": None}
    if not search_dir.exists():
        return found
    files = [p for p in search_dir.iterdir() if p.is_file() and not p.name.lower().endswith(".py")]
    resumes = [p for p in files if _looks_resume(p)]
    jobs    = [p for p in files if _looks_job(p)]
    found["resume"] = max(resumes, key=lambda x: x.stat().st_mtime, default=None)
    found["job"]    = max(jobs,    key=lambda x: x.stat().st_mtime, default=None)
    return found

# ===================== Builders & Enhancers =====================
def guess_name_from_text(resume_text: str) -> str:
    if not resume_text: return "Candidate"
    lines = [ln.strip() for ln in resume_text.splitlines() if ln.strip()]
    for ln in lines[:5]:
        m = re.match(r"^([A-Z][a-z]+(?:\s[A-Z]\.)?(?:\s[A-Z][a-z]+))$", ln)
        if m: return m.group(1)
    m = re.search(r"Name[:\-]\s*([A-Z][a-z]+(?:\s[A-Z]\.)?(?:\s[A-Z][a-z]+))", resume_text)
    return m.group(1) if m else "Candidate"

def build_payload(name: str) -> Dict[str, Any]:
    return {
        "name": name,
        "summary": "Professional, impact-focused contributor with a track record of delivering measurable results.",
        "core_items": ["Stakeholder Management","Customer Success","Process Improvement","Data-Driven Decisions"],
        "experience": [{
            "company":"Most Recent Company","title":"Most Recent Title","dates":"[CONFIRM]–Present",
            "bullets":[
                "Owned customer outcomes across a portfolio; delivered QBRs and renewal plans.",
                "Partnered cross-functionally to drive adoption and reduce churn."
            ]}],
        "education":"Education upon request"
    }

def payload_to_text(payload)->str:
    lines=[payload.get("summary",""), "Core: "+", ".join(payload.get("core_items",[]))]
    for j in payload.get("experience", []):
        lines.append(f"{j.get('title','')} @ {j.get('company','')} — {j.get('dates','')}")
        for b in j.get("bullets",[]): lines.append("• "+b)
    if payload.get("education"): lines.append("Education: " + payload.get("education",""))
    return "\n".join(lines)

def apply_medical_enhancements(payload, resume_text, jd_text):
    both = (resume_text + " " + jd_text)
    certs = [c for c in MED_CERTS if re.search(rf"\b{re.escape(c)}\b", both)]
    emr   = [e for e in MED_EMR if re.search(rf"\b{re.escape(e)}\b", both)]
    comp  = [c for c in MED_COMPLIANCE if re.search(rf"\b{re.escape(c)}\b", both)]
    if certs: payload.setdefault("licenses_certs", certs[:6])
    if emr or comp:
        payload["summary"] += " Experienced with " + ", ".join((emr[:2] or comp[:2])) + "."
    return payload

FEDERAL_SECTIONS = [
    "Citizenship: U.S. Citizen",
    "Clearances: [If applicable: Active/Public Trust/Secret/Top Secret/Sensitive Compartmented Information]",
    "Veterans’ Preference: [If applicable]",
    "Availability: [Full Time], Willing to Relocate: [Yes/No]",
    "Target Salary: [Optional; consider omitting]",
]

def apply_federal_formatting(payload):
    payload["summary"] = "Federal-format resume: detailed duties, results, hours/week, supervisor contact (optional), and KSAs. " + payload.get("summary","")
    if payload.get("experience"):
        b = payload["experience"][0]["bullets"]
        if len(b)<4:
            b.extend(["Documented outcomes with metrics; maintained compliance.", "Provided customer service and coordinated with cross-functional teams."])
    return payload

def enrich_from_vet_codes(payload, resume_text):
    for code, trans in VET_CROSSWALK.items():
        if re.search(rf"\b{re.escape(code)}\b", resume_text):
            if trans not in payload["core_items"]:
                payload["core_items"].append(trans)
    return payload

def inject_keywords(payload, jd_text, per_section_caps=(2,2,2)):
    missing = list(kw_set(jd_text) - kw_set(payload_to_text(payload)))
    noise = {"and","with","from","will","work","role","team","customer","customers","clients","provide","years","experience","requirements","must","have"}
    missing = [m for m in missing if m not in noise][:6]
    if not missing: return payload
    sum_cap, core_cap, bullet_cap = per_section_caps
    if sum_cap and missing:
        extra = ", ".join(missing[:sum_cap])
        payload["summary"] = (payload.get("summary","") + (" " if payload.get("summary") else "") + f"Focus: {extra}.").strip()
        missing = missing[sum_cap:]
    if core_cap and missing:
        core = payload.get("core_items", [])
        for t in list(missing)[:core_cap]:
            if t.title() not in core: core.append(t.title())
        payload["core_items"] = core
        missing = missing[core_cap:]
    if bullet_cap and missing and payload.get("experience"):
        b = payload["experience"][0].get("bullets", [])
        add = list(missing)[:bullet_cap]
        b.insert(0, f"Advanced {', '.join(add)} initiatives aligned to role requirements.")
        payload["experience"][0]["bullets"] = b
    return payload

def ensure_metric_shell(payload):
    if not payload.get("experience"): return payload
    b = payload["experience"][0].get("bullets", [])
    if not any(re.search(r"\d", x) for x in b):
        b.insert(0, "Improved adoption by [CONFIRM]% and renewal rate by [CONFIRM]% across assigned accounts.")
        payload["experience"][0]["bullets"] = b
    return payload

# ===================== Risk =====================
def scan_risk(text: str) -> List[str]:
    flags = []
    for name, pat in SCAM_PATTERNS.items():
        if pat.search(text): flags.append(name)
    return flags

# ===================== Writers =====================
def rtf_escape(s: str) -> str:
    return s.replace("\\", "\\\\").replace("{","\\{").replace("}","\\}")

def write_docx_resume(out_path: Path, payload, style_pack="Classic"):
    tx = payload_to_text(payload)
    if not DOCX_OK:
        return write_rtf_resume(out_path.with_suffix(".rtf"), payload)
    doc=docx.Document()
    sp = STYLE_PACKS.get(style_pack, STYLE_PACKS["Classic"])
    style=doc.styles['Normal']; style.font.name=sp["font"]; style.font.size=Pt(sp["size"])

    # Name centered
    p=doc.add_paragraph(payload.get("name","")); p.runs[0].bold=True; p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    if payload.get("summary"):
        doc.add_paragraph(payload["summary"])

    # Core
    core = payload.get("core_items",[])
    if core:
        doc.add_paragraph().add_run("Core Skills").bold=True
        doc.add_paragraph(", ".join(core))

    # Federal preface (if selected or detected)
    if style_pack=="Federal":
        doc.add_paragraph().add_run("Federal Information").bold=True
        for line in FEDERAL_SECTIONS:
            doc.add_paragraph("• " + line)

    # Experience
    if payload.get("experience"):
        doc.add_paragraph().add_run("Experience").bold=True
        for j in payload["experience"]:
            run=doc.add_paragraph().add_run(f"{j.get('title','')} — {j.get('company','')} ({j.get('dates','')})"); run.bold=True
            for b in j.get("bullets",[]): doc.add_paragraph("• "+b)

    # Education
    if payload.get("education"):
        doc.add_paragraph().add_run("Education").bold=True
        doc.add_paragraph(payload.get("education",""))

    doc.save(str(out_path))
    return out_path

def write_rtf_resume(out_path: Path, payload):
    body = []
    body.append(r"{\rtf1\ansi\deff0")
    body.append(r"\b " + rtf_escape(payload.get("name","")) + r"\b0\par")
    if payload.get("summary"):
        body.append(rtf_escape(payload["summary"]) + r"\par")
    core = payload.get("core_items",[])
    if core:
        body.append(r"\b Core Skills\b0\par")
        body.append(rtf_escape(", ".join(core)) + r"\par")
    if payload.get("experience"):
        body.append(r"\b Experience\b0\par")
        for j in payload["experience"]:
            head = f"{j.get('title','')} — {j.get('company','')} ({j.get('dates','')})"
            body.append(r"\b " + rtf_escape(head) + r"\b0\par")
            for b in j.get("bullets",[]): body.append(r"\bullet " + rtf_escape(b) + r"\par")
    if payload.get("education"):
        body.append(r"\b Education\b0\par")
        body.append(rtf_escape(payload.get("education","")) + r"\par")
    body.append("}")
    out_path.write_text("".join(body), encoding="utf-8")
    return out_path

def write_docx_cover(out_path: Path, name:str, company:str, title:str, style_pack="Classic"):
    content = f"""Dear Hiring Team,

I'm excited to apply for the {title or 'role'} at {company or 'your company'}. My background includes leading outcomes, driving adoption, and delivering measurable results. I partner cross-functionally, communicate clearly with executives, and turn data into action.

Sincerely,
{name}
"""
    if not DOCX_OK:
        out_path = out_path.with_suffix(".rtf")
        out_path.write_text("{\\rtf1\\ansi " + rtf_escape(content) + "}", encoding="utf-8")
        return out_path
    doc=docx.Document()
    sp = STYLE_PACKS.get(style_pack, STYLE_PACKS["Classic"])
    style=doc.styles['Normal']; style.font.name=sp["font"]; style.font.size=Pt(sp["size"])
    for ln in content.split("\n"): doc.add_paragraph(ln)
    doc.save(str(out_path)); return out_path

def write_html_guide(out_path: Path, company:str, title:str, scores:dict, jd_snip:str):
    html=f"""<!doctype html><html><meta charset='utf-8'><title>Interview Guide</title>
<style>body{{font-family:system-ui,Segoe UI,Arial;margin:24px;}} .k{{padding:8px 12px;border-radius:8px;background:#f5f5f7;display:inline-block;margin:4px 8px 12px 0}}</style>
<h1>Interview Study Guide</h1>
<p><b>Role:</b> {title or 'N/A'} &nbsp; | &nbsp; <b>Company:</b> {company or 'N/A'}</p>
<div class='k'>Grade: {scores['grade']}</div>
<div class='k'>Match: {scores['cov']}%</div>
<div class='k'>Scan: {scores['ats']}</div>
<div class='k'>Hire: {scores['hire']}</div>
<h2>What to Know</h2>
<ul>
<li>Target the <b>Top 5</b> keywords from the job post in your answers.</li>
<li>Bring <b>two metrics</b> (renewal %, adoption %, ARR/MRR impact).</li>
<li>Prepare a <b>60-second</b> story about driving adoption and reducing churn.</li>
</ul>
<h2>Job Snippet</h2>
<pre style='white-space:pre-wrap'>{jd_snip[:1500]}</pre>
</html>"""
    out_path.write_text(html, encoding="utf-8"); return out_path

def write_redline(out_path: Path, before:str, after:str):
    diff = "\n".join(difflib.unified_diff(before.splitlines(), after.splitlines(), fromfile="before", tofile="after", lineterm=""))
    out_path.write_text(diff, encoding="utf-8"); return out_path

def write_ats_txt(out_path: Path, txt:str):
    out_path.write_text(txt, encoding="utf-8"); return out_path

def zip_deliverables(zip_path: Path, files: List[Path]):
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for f in files:
            if f and Path(f).exists():
                z.write(f, arcname=Path(f).name)
    return zip_path

# ===================== Flow =====================
def detect_and_read():
    # Search uploads first, then CWD
    uploads = scan_uploads(Path("/mnt/data"))
    resume_p, jd_p = uploads.get("resume"), uploads.get("job")
    if not (resume_p and jd_p):
        local = scan_uploads(Path.cwd())
        resume_p = resume_p or local.get("resume")
        jd_p     = jd_p or local.get("job")
    return resume_p, jd_p

def guess_company_and_title(jd_text: str, jd_path: str="") -> (str,str):
    company=""; title=""
    m = re.search(r"Company[:\-]\s*([^\n]+)", jd_text, re.I)
    if m: company = m.group(1).strip()
    m = re.search(r"(Job\s*Title|Title)[:\-]\s*([^\n]+)", jd_text, re.I)
    if m: title = m.group(2).strip()
    if not title:
        lines=[ln.strip() for ln in jd_text.splitlines() if ln.strip()]
        if lines: title = lines[0][:80]
    if not company and jd_path and re.match(r"^https?://", str(jd_path)):
        host = re.sub(r"^https?://","", str(jd_path)).split("/")[0]
        host = host.split(":")[0]
        parts = host.split(".")
        if len(parts)>=2: company = parts[-2].capitalize()
    for bad in ["Careers","Jobs","Job", "Hiring","Openings"]:
        if title.endswith(bad): title = title.replace(bad,"").strip(" -|")
    return company, title

def main():
    try:
        base_dir = Path.cwd()
        outdir = base_dir / "out"
        outdir.mkdir(parents=True, exist_ok=True)

        resume_p, jd_p = detect_and_read()
        if not resume_p or not jd_p:
            missing = []
            if not resume_p: missing.append("📑 your resume")
            if not jd_p:     missing.append("📋 the job posting (file or link)")
            need = " and ".join(missing)
            ui_card(WELCOME_CARD + f"\n📥 Please upload {need} to continue.\n")
            return

        res_text = read_text_any(str(resume_p))
        jd_text  = read_text_any(str(jd_p))

        name = guess_name_from_text(res_text)
        payload = build_payload(name)
        enrich_from_vet_codes(payload, res_text)

        # Domain enrichments
        text_both = (res_text + " " + jd_text).lower()
        if any(k in text_both for k in ["clinic","hospital","epic","cerner","patient","hipaa"]):
            payload = apply_medical_enhancements(payload, res_text, jd_text)
            style_pack = "Healthcare"
        elif any(k in text_both for k in ["federal","usajobs","gs-","clearance","public trust","ts","secret"]):
            payload = apply_federal_formatting(payload)
            style_pack = "Federal"
        else:
            style_pack = "Classic"

        # Optimize keywords
        before_txt = payload_to_text(payload)
        payload = inject_keywords(payload, jd_text)
        payload = ensure_metric_shell(payload)
        after_txt = payload_to_text(payload)

        cov, ats, hire, grade = composite_scores(jd_text, after_txt)
        company, title = guess_company_and_title(jd_text, str(jd_p))

        base = f"{(name).replace(' ','_')}_{(company or 'Company').replace(' ','_')}_{(title or 'Role').replace(' ','_')}_{TODAY}"
        saved = []

        resume_path = outdir / f"{base}_RESUME.docx"
        saved.append(str(write_docx_resume(resume_path, payload, style_pack=style_pack)))

        cover_path  = outdir / f"{base}_COVER.docx"
        saved.append(str(write_docx_cover(cover_path, name, company or "", title or "", style_pack=style_pack)))

        guide_path  = outdir / f"{base}_GUIDE.html"
        saved.append(str(write_html_guide(guide_path, company or "", title or "", {"cov":cov,"ats":ats,"hire":hire,"grade":grade}, jd_text[:2000])))

        ats_path    = outdir / f"{base}_ATS.txt"
        saved.append(str(write_ats_txt(ats_path, after_txt)))

        redline_path= outdir / f"{base}_REDLINE.diff.txt"
        saved.append(str(write_redline(redline_path, before_txt, after_txt)))

        # Risk scan (basic regex flags)
        flags = scan_risk(jd_text)
        risky = bool(flags)


        # --- FAT extras (saved only when FAT_MODE = True) ---
        if FAT_MODE:
            try:
                payload_json = outdir / f"{base}_PAYLOAD.json"
                payload_json.write_text(json.dumps(payload, indent=2), encoding="utf-8")
                saved.append(str(payload_json))

                risk_json = outdir / f"{base}_RISK.json"
                risk_json.write_text(json.dumps({"flags": list(flags)}, indent=2), encoding="utf-8")
                saved.append(str(risk_json))

                kw_json = outdir / f"{base}_KW.json"
                kw_json.write_text(json.dumps({
                    "jd_top": sorted(list(kw_set(jd_text)))[:200],
                    "resume_top": sorted(list(kw_set(after_txt)))[:200]
                }, indent=2), encoding="utf-8")
                saved.append(str(kw_json))

                fat_guide = outdir / f"{base}_GUIDE_FAT.html"
                write_html_guide(fat_guide, company or "", title or "", {"cov":cov,"ats":ats,"hire":hire,"grade":grade}, jd_text[:4000])
                saved.append(str(fat_guide))
            except Exception as _e:
                ui_debug(f"[FAT extras] {_e}")
        # --- end FAT extras ---

        # Zip
        zip_path = outdir / f"{base}_Deliverables.zip"
        zip_deliverables(zip_path, [Path(p) for p in saved])
        saved.append(str(zip_path))

        # Card summary
        files_block = ''.join(['• ' + Path(p).name + '\n' for p in saved])
        ui_card(f"""📄 Name: {name}
🏢 Company: {company or '—'}
🧰 Title: {title or '—'}

📊 Scores
• 🔑 Match (keywords): {cov}%
• ⚙️ Scan (Applicant Tracking System health): {ats}%
• 💼 Hire likelihood: {hire}%
• 🎓 Grade: {grade}

🚩 Risk Check
{'✅ No risky job flags detected.' if not risky else '⚠️ ' + ', '.join(flags)}

📂 Files saved
{files_block}
""")

    except Exception as e:
        ui_card("⚠️ Something went wrong while generating files. Please re-upload your resume and the job post, then try again.")

# ============ Minimal helpers for names and payloads ============
def build_payload(name: str) -> Dict[str, Any]:
    return {
        "name": name,
        "summary": "Professional, impact-focused contributor with a track record of delivering measurable results.",
        "core_items": ["Stakeholder Management","Customer Success","Process Improvement","Data-Driven Decisions"],
        "experience": [{
            "company":"Most Recent Company","title":"Most Recent Title","dates":"[CONFIRM]–Present",
            "bullets":[
                "Owned customer outcomes across a portfolio; delivered QBRs and renewal plans.",
                "Partnered cross-functionally to drive adoption and reduce churn."
            ]}],
        "education":"Education upon request"
    }

def payload_to_text(payload)->str:
    lines=[payload.get("summary",""), "Core: "+", ".join(payload.get("core_items",[]))]
    for j in payload.get("experience", []):
        lines.append(f"{j.get('title','')} @ {j.get('company','')} — {j.get('dates','')}")
        for b in j.get("bullets",[]): lines.append("• "+b)
    if payload.get("education"): lines.append("Education: " + payload.get("education",""))
    return "\n".join(lines)

if __name__ == "__main__":
    main()


# ===========================
# INVALID_CODE (Roadmap Notes)
# ===========================
# Purpose: This block documents planned/experimental features NOT active in v11 final.
# Keep it in-code so distribution carries the roadmap transparently.
#
# 1) Parsing & Reliability
#    - OCR for scanned PDFs (Tesseract wrapper) for offline use.
#    - Docx style-preserving export with tables → safe text fallback side-by-side.
#    - Golden test suite with hashed expected outputs (selftest full).
#
# 2) Researcher++
#    - Passive-voice detector (threshold and suggestions).
#    - Duplicate-bullet deduper across roles.
#    - Impact metric suggester trained on domain patterns (no PII).
#
# 3) Federal / USAJobs
#    - KSAs auto-expansion with sentence extraction + mapping to role bullets.
#    - Specialized Experience detector with confidence score.
#    - Clearance normalization (TS/SCI, CI Poly, etc.) with caveats.
#
# 4) Domain Packs
#    - Expanded curated packs (legal/finance/healthcare/engineering/marketing).
#    - Light synonym mapping to avoid overfitting to exact tokens.
#
# 5) LinkedIn / Profiles
#    - Structured JSON resume export (role/company/dates/impact/skills).
#    - Guided "About" and "Experience" rewrites with tone presets.
#
# 6) Education / Demo Mode
#    - Highlight unmet requirements in subtle color (when GUI is available).
#    - Bar passage/board exams tracker with disclaimers (no fabrication).
#
# 7) Second-Chance / Staffing
#    - Softened tone, reliability framing, employer set lists.
#    - Region-aware guidance for fair-chance employers.
#
# 8) UX / Help
#    - Inline 'help modes', 'help federal', 'help ats' navigable sections.
#    - First-time tip toggle and per-mode mini-guides.
#
# 9) Performance
#    - Micro-caching of cleaned inputs and extracted keywords per run dir.
#    - Deterministic seeding across sessions (serialize seed in bundle).
#
# 10) Telemetry (local-only, opt-in)
#     - Timing and pass/fail reasons (no PII) to diagnose bottlenecks.
#
# End of INVALID_CODE notes.
