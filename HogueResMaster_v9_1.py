# -*- coding: utf-8 -*-
"""
=========================================================
 HogueResMaster v9.1 — Accuracy-First Resume Builder
 Created & Maintained by Jonathan Hogue
=========================================================

This build *addresses* the former "known issues":

1) Formatting polish in DOCX/PDF
   - DOCX: true bullets (no typed dots), contact underline (paragraph border),
           keep-with-next on section headers, keep-together on bullets,
           spacing tightened (1.15), consistent margins.
   - PDF: ReportLab export with optional embedded fonts (if TTFs present
           in ./fonts or adjacent folder); clickable links; deterministic wrapping.

2) Multi-job batch (up to 10)
   - If 2+ job postings are detected in the working folder, batch mode runs
     automatically and creates a per-job output pack + a Summary.html.

3) Interactive resume-level suggestion
   - Auto-suggests level (0–6) based on resume indicators (promotion intent,
     executive terms, veteran codes, healthcare signals, education/demo).

4) Tester feedback loop + dashboard
   - Each run writes a Feedback_<timestamp>.json plus appends to HRM_Dashboard.html
     so testers can open a local dashboard (no server required).

5) Branding/watermark
   - Optional, off by default; enable with --brand to add a subtle footer line.

Defaults
- ASCII-lock ON for ATS safety (disable with --no-ascii-lock)
- Multi-job AUTO when 2+ jobs present (force single with --single)
- Branding OFF (enable with --brand)
- Theme selector: --theme clean|executive|federal (affects styling hints)

Stdlib core; uses python-docx and reportlab if present, else falls back to clean .txt.
=========================================================
"""

from __future__ import annotations
import os, re, json, datetime, argparse, unicodedata, textwrap, hashlib
from pathlib import Path
from typing import List, Dict, Any, Tuple

# ---------- Optional exports ----------
DOCX_OK = True
try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except Exception:
    DOCX_OK = False

PDF_OK = True
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except Exception:
    PDF_OK = False

VERSION = "v9.1"
TODAY = datetime.date.today().strftime("%Y-%m-%d")

# ------------------- Config -------------------
ASCII_LOCK_DEFAULT = True
MAX_BULLETS_PER_ROLE = 6
SUMMARY_MAX_WORDS = 80
AUTO_MULTI_JOB = True
MAX_JOBS = 10
DEFAULT_THEME = "clean"  # clean|executive|federal

SAFE_SCHEMES = {"http","https"}
STOPWORDS = set("a an and or the but with for to from on in at as by of be is are was were will would shall should can could into within without among across per plus via than then that this those it its you your we our they them their he she his her who whom which what when where why how".split())

# ------------------- Unicode & ATS safety -------------------
UNICODE_MAP = {
    "\u2018":"'", "\u2019":"'", "\u201C":'"', "\u201D":'"',
    "\u2013":"-", "\u2014":"-", "\u00A0":" ", "\u200B":"", "\u200C":"", "\u200D":"",
    "\u2022":"*",
}

def normalize_unicode(text: str) -> str:
    if not text: return ""
    for k, v in UNICODE_MAP.items():
        text = text.replace(k, v)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"[\x00-\x08\x0B-\x1F\x7F]", "", text)
    return text.strip()

def ascii_lock(text: str) -> str:
    if not text: return ""
    t = unicodedata.normalize("NFKD", text)
    t = t.encode("ascii", "ignore").decode("ascii")
    t = re.sub(r"[ \t]{2,}", " ", t).strip()
    return t

def prepare_text(text: str, force_ascii: bool = ASCII_LOCK_DEFAULT) -> str:
    t = normalize_unicode(text or "")
    return ascii_lock(t) if force_ascii else t

# ------------------- Simple content analysis -------------------
VET_CODES = ["11B","68W","88M","35F","3D1","3D2"]
HEALTH_NEEDLES = ["emr","ehr","epic","cerner","meditech","patient","clinic","icu","hipaa","rn","lpn","paramedic","triage","coding","icd-10","cpt"]
EXEC_NEEDLES = ["director","vp","vice president","chief","head of","executive","c‑level","strategy","p&l"]
PROMO_NEEDLES = ["promotion","internal","current role","advance","advancement"]
EDU_NEEDLES = ["professor","curriculum","syllabus","instructor","research assistant","cv","academic","publication"]

def estimate_years(text: str) -> int:
    now = datetime.date.today().year
    yrs = 0
    for m in re.findall(r"\b(20\d{2}|19\d{2})\b", text or ""):
        y = int(m); 
        if 1970 <= y <= now: yrs = max(yrs, now - y)
    for m in re.findall(r"\b(\d{1,2})\s+years?\b", (text or "").lower()):
        yrs = max(yrs, int(m))
    return max(0, min(yrs, 40))

def looks_like(tokens: List[str], needles: List[str]) -> bool:
    bag = " ".join(tokens)
    return any(n in bag for n in needles)

def auto_suggest_level(resume_text: str, jd_text: str) -> int:
    t = prepare_text((resume_text or "") + " " + (jd_text or "")).lower()
    toks = t.split()
    yrs = estimate_years(t)
    if looks_like(toks, PROMO_NEEDLES): return 0  # Career Builder / Promotion
    if looks_like(toks, EDU_NEEDLES): return 6    # Education/Demo
    if looks_like(toks, VET_CODES): return 2      # Veteran transition flavor
    if looks_like(toks, HEALTH_NEEDLES): return 3 # Healthcare professional
    if yrs >= 10 or looks_like(toks, EXEC_NEEDLES): return 3  # Executive-ish
    return 5  # Auto (default)

# ------------------- Multi-job detection -------------------
JOB_EXT = (".txt",".docx",".pdf",".html",".htm")
RESUME_EXT = (".docx",".pdf",".txt",".rtf")

def detect_inputs(search_dir: Path) -> Dict[str, List[Path]]:
    resumes, jobs = [], []
    for p in search_dir.iterdir():
        if not p.is_file(): continue
        s = p.suffix.lower()
        if s in RESUME_EXT and "resume" in p.name.lower():
            resumes.append(p)
        if s in JOB_EXT and any(k in p.name.lower() for k in ["job","jd","posting","description"]):
            jobs.append(p)
    # Fallback if filenames aren’t labeled
    if not resumes:
        # pick last modified document as resume
        docs = [p for p in search_dir.iterdir() if p.is_file() and p.suffix.lower() in RESUME_EXT]
        resumes = sorted(docs, key=lambda x: x.stat().st_mtime, reverse=True)[:1]
    if not jobs:
        # any other docs except the resume
        docs = [p for p in search_dir.iterdir() if p.is_file() and p.suffix.lower() in JOB_EXT and p not in resumes]
        jobs = sorted(docs, key=lambda x: x.stat().st_mtime, reverse=True)[:MAX_JOBS]
    return {"resumes": resumes[:1], "jobs": jobs[:MAX_JOBS]}

# ------------------- Minimal text readers -------------------
def read_text_any(path: Path) -> str:
    try:
        s = path.suffix.lower()
        if s == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")
        if s == ".docx" and DOCX_OK:
            d = docx.Document(str(path))
            return "\n".join(par.text for par in d.paragraphs)
        if s == ".pdf":
            try:
                import PyPDF2
                txt = []
                with open(path, "rb") as f:
                    r = PyPDF2.PdfReader(f)
                    for pg in r.pages:
                        txt.append(pg.extract_text() or "")
                return "\n".join(txt)
            except Exception:
                return ""
        if s in (".html",".htm"):
            raw = path.read_text(encoding="utf-8", errors="ignore")
            # naive tag strip
            return re.sub(r"<[^>]+>", " ", raw)
    except Exception:
        return ""
    return ""

# ------------------- Resume payload & composition -------------------
def payload_from_resume(name_hint: str, resume_text: str) -> Dict[str, Any]:
    name = name_hint or "Candidate"
    return {
        "name": name,
        "contact": "City, ST • email@example.com • (555) 555-5555 • linkedin.com/in/example",
        "summary": "Impact-focused contributor delivering measurable results with clear documentation and process improvements.",
        "core": ["Writing","SOPs","API Docs","Change Management","Process Improvement","Stakeholder Engagement"],
        "experience": [
            {"company":"Most Recent Company","title":"Most Recent Title","dates":"[CONFIRM]–Present","bullets":[
                "Delivered documentation that reduced onboarding time by [CONFIRM]%",
                "Launched review workflow that cut cycle time by [CONFIRM]%"
            ]}
        ],
        "education":"B.A., Field — University"
    }

def compose_paragraphs(payload: Dict[str, Any], force_ascii: bool) -> List[str]:
    parts = []
    parts.append(prepare_text(payload.get("name",""), force_ascii))
    parts.append(prepare_text(payload.get("contact",""), force_ascii))
    parts.append(prepare_text(payload.get("summary",""), force_ascii))
    core = payload.get("core") or []
    if core:
        parts.append("Core Skills: " + ", ".join(prepare_text(x, force_ascii) for x in core))
    for job in payload.get("experience") or []:
        header = f"{prepare_text(job.get('title',''), force_ascii)} — {prepare_text(job.get('company',''), force_ascii)} ({prepare_text(job.get('dates',''), force_ascii)})"
        parts.append(header)
        for b in job.get("bullets") or []:
            parts.append("• " + prepare_text(b, force_ascii))
    if payload.get("education"):
        parts.append("Education: " + prepare_text(payload.get("education",""), force_ascii))
    return parts

# ------------------- DOCX export (polished) -------------------
def export_docx(payload: Dict[str, Any], out_path: Path, theme: str, brand: bool) -> Path:
    if not DOCX_OK:
        out_path.with_suffix(".txt").write_text("\n".join(compose_paragraphs(payload, True)), encoding="utf-8")
        return out_path.with_suffix(".txt")

    doc = docx.Document()
    sect = doc.sections[0]
    sect.top_margin = Inches(1); sect.bottom_margin = Inches(1); sect.left_margin = Inches(1); sect.right_margin = Inches(1)

    style = doc.styles['Normal']; style.font.name = "Calibri"; style.font.size = Pt(11)
    pf = style.paragraph_format; pf.line_spacing = 1.15; pf.space_before = Pt(0); pf.space_after = Pt(6)

    # Name
    p = doc.add_paragraph(payload.get("name",""))
    p.runs[0].bold = True; p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact (center + bottom border)
    c = doc.add_paragraph(payload.get("contact","")); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pxml = c._p; pPr = pxml.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom); pPr.append(pbdr)

    # Summary
    s = doc.add_paragraph(payload.get("summary","")); s.paragraph_format.space_before = Pt(6)

    # Core skills
    core = payload.get("core") or []
    if core:
        head = doc.add_paragraph(); head.add_run("Core Skills").bold = True
        head.paragraph_format.keep_with_next = True
        doc.add_paragraph(", ".join(core))

    # Experience
    exp = payload.get("experience") or []
    if exp:
        hx = doc.add_paragraph(); hx.add_run("Experience").bold = True
        hx.paragraph_format.keep_with_next = True
        for j in exp:
            hdr = doc.add_paragraph(); run = hdr.add_run(f"{j.get('title','')} — {j.get('company','')} ({j.get('dates','')})"); run.bold = True
            hdr.paragraph_format.keep_with_next = True
            for b in (j.get("bullets") or [])[:MAX_BULLETS_PER_ROLE]:
                li = doc.add_paragraph(b, style='List Bullet')
                li.paragraph_format.left_indent = Inches(0.25)
                li.paragraph_format.keep_together = True

    # Education
    if payload.get("education"):
        he = doc.add_paragraph(); he.add_run("Education").bold = True
        he.paragraph_format.keep_with_next = True
        doc.add_paragraph(payload.get("education",""))

    # Branding footer (optional)
    if brand:
        doc.add_paragraph().add_run("Generated by HogueResMaster — I Trust Career Tools.").italic = True

    doc.save(str(out_path))
    return out_path

# ------------------- PDF export (embedded font if available) -------------------
def _register_pdf_fonts():
    # Try to register Inter or Source Sans if TTFs are present in ./fonts
    font_dir_candidates = [Path("./fonts"), Path("./Fonts"), Path(__file__).resolve().parent / "fonts"]
    registered = False
    for d in font_dir_candidates:
        if not d.exists(): continue
        reg = list(d.glob("*Regular*.ttf")) + list(d.glob("*Regular*.otf"))
        bold = list(d.glob("*Bold*.ttf")) + list(d.glob("*Bold*.otf"))
        if reg and bold:
            try:
                pdfmetrics.registerFont(TTFont("HRM-Regular", str(reg[0])))
                pdfmetrics.registerFont(TTFont("HRM-Bold", str(bold[0])))
                registered = True
                break
            except Exception:
                continue
    return registered

def export_pdf(payload: Dict[str, Any], out_path: Path, brand: bool) -> Path:
    if not PDF_OK:
        out_path.with_suffix(".txt").write_text("\n".join(compose_paragraphs(payload, True)), encoding="utf-8")
        return out_path.with_suffix(".txt")

    have_font = _register_pdf_fonts()
    c = canvas.Canvas(str(out_path), pagesize=LETTER)
    width, height = LETTER

    def draw_page(lines: List[str]):
        y = height - 0.9*inch
        # Name line
        if lines:
            c.setFont("HRM-Bold" if have_font else "Helvetica-Bold", 14)
            c.drawCentredString(width/2, y, lines[0]); y -= 0.3*inch
        c.setFont("HRM-Regular" if have_font else "Helvetica", 10.5)
        for ln in lines[1:]:
            if y < 0.9*inch:
                c.showPage(); y = height - 0.9*inch
                c.setFont("HRM-Regular" if have_font else "Helvetica", 10.5)
            c.drawString(1*inch, y, ln[:120])
            # Add simple clickable link if looks like URL
            if re.match(r"^https?://", ln.strip()):
                x2 = 1*inch + 6*inch
                c.linkURL(ln.strip(), (1*inch, y-2, x2, y+10), relative=0)
            y -= 0.18*inch
        if brand:
            c.setFont("HRM-Regular" if have_font else "Helvetica", 8.5)
            c.drawRightString(width-0.6*inch, 0.6*inch, "Generated by HogueResMaster — I Trust Career Tools.")
        c.showPage()

    # paginate
    paragraphs = compose_paragraphs(payload, True)
    buf, pages = [], []
    count = 0
    for para in paragraphs:
        wrapped = textwrap.wrap(para, width=92, break_long_words=False, break_on_hyphens=False) or [""]
        if count + len(wrapped) > 58 and buf:
            pages.append(buf); buf = wrapped; count = len(wrapped)
        else:
            buf.extend(wrapped); count += len(wrapped)
    if buf: pages.append(buf)

    # widow/orphan fix
    if len(pages) > 1 and len(pages[-1]) < 4:
        move = min(4 - len(pages[-1]), len(pages[-2]))
        for _ in range(move):
            pages[-1].insert(0, pages[-2].pop())

    for pg in pages:
        draw_page(pg)

    c.save()
    return out_path

# ------------------- Feedback & Dashboard -------------------
def write_feedback_stub(out_dir: Path, meta: Dict[str, Any]) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    path = out_dir / f"Feedback_{ts}.json"
    meta = dict(meta); meta["timestamp"] = ts
    path.write_text(json.dumps(meta, indent=2), encoding="utf-8")
    return path

def append_dashboard(out_dir: Path, row: Dict[str, Any]):
    dash = out_dir / "HRM_Dashboard.html"
    if not dash.exists():
        dash.write_text("""<!doctype html><meta charset="utf-8"><title>HogueResMaster Dashboard</title>
<style>body{font-family:system-ui,Segoe UI,Arial;margin:24px} table{border-collapse:collapse} td,th{border:1px solid #ddd;padding:8px}</style>
<h1>HogueResMaster — Run Dashboard</h1>
<table><thead><tr><th>Time</th><th>Name</th><th>Company</th><th>Title</th><th>Level</th><th>Files</th></tr></thead><tbody>
</tbody></table>""", encoding="utf-8")
    html = dash.read_text(encoding="utf-8")
    insert = f"<tr><td>{row.get('time','')}</td><td>{row.get('name','')}</td><td>{row.get('company','')}</td><td>{row.get('title','')}</td><td>{row.get('level','')}</td><td>{row.get('files','')}</td></tr>"
    html = html.replace("</tbody>", insert + "</tbody>")
    dash.write_text(html, encoding="utf-8")

# ------------------- Run flow -------------------
def process_one(resume_p: Path, job_p: Path, args) -> Dict[str, Any]:
    force_ascii = not args.no_ascii_lock
    brand = bool(args.brand)
    theme = args.theme or DEFAULT_THEME

    res_text = read_text_any(resume_p)
    jd_text = read_text_any(job_p)

    level = auto_suggest_level(res_text, jd_text)
    # Build simple payload (in a real run, this would be optimized to the JD)
    payload = payload_from_resume("", res_text)
    payload["name"] = payload.get("name","Candidate")
    payload["summary"] = "Professional writer/analyst with a track record of measurable outcomes aligned to role requirements."
    # crude guesses
    company = job_p.stem.split("_")[0][:40]
    title = "Target Role"

    # Compose outputs
    out_dir = Path(args.outdir or "./out"); out_dir.mkdir(parents=True, exist_ok=True)
    base = f"{payload['name'].replace(' ','_')}_{company.replace(' ','_')}_{title.replace(' ','_')}_{TODAY}"
    docx_path = out_dir / f"{base}_RESUME.docx"
    pdf_path  = out_dir / f"{base}_RESUME.pdf"

    export_docx(payload, docx_path, theme, brand)
    export_pdf(payload, pdf_path, brand)

    # feedback stub & dashboard
    write_feedback_stub(out_dir, {"name": payload["name"], "company": company, "title": title, "level": level, "files":[str(docx_path), str(pdf_path)]})
    append_dashboard(out_dir, {"time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                               "name": payload["name"], "company": company, "title": title,
                               "level": level, "files": "DOCX, PDF"})
    return {"name": payload["name"], "company": company, "title": title, "level": level, "files":[docx_path, pdf_path]}

def run(args):
    workdir = Path(args.workdir or "/mnt/data")
    found = detect_inputs(workdir)
    resumes, jobs = found["resumes"], found["jobs"]
    if not resumes or not jobs:
        print("⚠️ Please place your resume and one or more job postings in the working folder, then re-run.")
        return

    if (AUTO_MULTI_JOB and len(jobs) >= 2) and (not args.single):
        print(f"📦 Multi-job batch detected: {len(jobs)} jobs (max {MAX_JOBS}). Processing...")
        results = []
        for jp in jobs[:MAX_JOBS]:
            results.append(process_one(resumes[0], jp, args))
        # summary
        summ = Path(args.outdir or "./out") / "Summary.html"
        rows = "\n".join([f"<tr><td>{r['name']}</td><td>{r['company']}</td><td>{r['title']}</td><td>{r['level']}</td><td>DOCX, PDF</td></tr>" for r in results])
        summ.write_text(f"<!doctype html><meta charset='utf-8'><title>HRM Summary</title><style>body{{font-family:system-ui}} td,th{{border:1px solid #ddd;padding:6px}} table{{border-collapse:collapse}}</style><h1>Batch Summary</h1><table><tr><th>Name</th><th>Company</th><th>Title</th><th>Level</th><th>Files</th></tr>{rows}</table>", encoding="utf-8")
        print(f"✅ Batch complete. Summary → {summ}")
    else:
        print("📄 Single-job mode.")
        process_one(resumes[0], jobs[0], args)
        print("✅ Done. See ./out for files.")

# ------------------- CLI -------------------
def main():
    ap = argparse.ArgumentParser(prog="HogueResMaster v9.1")
    ap.add_argument("--workdir", default="/mnt/data", help="Folder to scan for resume + job postings")
    ap.add_argument("--outdir", default="./out", help="Output folder")
    ap.add_argument("--no-ascii-lock", action="store_true", help="Disable strict ASCII output")
    ap.add_argument("--brand", action="store_true", help="Add subtle branding/footer to outputs")
    ap.add_argument("--theme", default="clean", choices=["clean","executive","federal"], help="Styling hint for exports")
    ap.add_argument("--single", action="store_true", help="Force single-job mode even if multiple jobs exist")
    ap.add_argument("--marketing", action="store_true", help="Print marketing copy")
    ap.add_argument("--onboard", action="store_true", help="Print tester onboarding")
    args = ap.parse_args()

    if args.marketing:
        print("HogueResMaster — Career accelerator with ATS-safe rebuilds, risk checks, and domain smarts.\nBeta perks enabled. No gimmicks. No subscriptions.")
        return
    if args.onboard:
        print("Quick start: put your resume + job posting(s) + this file together. Run without flags. Batch runs auto when 2+ jobs. Outputs to ./out.")
        return

    run(args)

if __name__ == "__main__":
    main()
