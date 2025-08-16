
# -*- coding: utf-8 -*-
"""
=========================================================
 HogueResMaster – v9.2.2 (Mobile-first, ATS-safe, Pro Output)
 Created & Coded by Jonathan Hogue
=========================================================

This build implements:
  • Issue 1: RTF fallback for Resume/Cover when python-docx is unavailable (no more .txt)
  • Issue 2: Real PDF Interview Guide when ReportLab is available (HTML kept as fallback)
  • Issue 3: Update checker is opt-in via --check_updates yes|no (default: no)

Mission (external)
---------------------------------------------------------
Level the playing field. Every job seeker deserves a resume that
gets noticed — professional, fast, and safe.

How to run (mobile-first)
---------------------------------------------------------
1) In one chat (ChatGPT, Gemini, or Copilot), upload:
   • Your resume (Word/PDF/screenshot)
   • The job posting (file/screenshot/link)
   • This file: HogueResMaster_v9_2_2.py
2) Type exactly: Run HogueResMaster
3) You’ll get: Resume (DOCX or RTF*), Cover Letter (DOCX or RTF*), Interview Guide (HTML + PDF*),
   ATS-safe Text, Redline (diff), Manifest JSON
   * RTF is produced when python-docx is not available. PDF is produced when ReportLab is installed.

Notes
---------------------------------------------------------
• Works even if optional libs are missing (graceful fallbacks without losing professional formatting).
• Acronyms expand on first mention (e.g., Applicant Tracking System (ATS)).
• Update checker only runs when you pass --check_updates yes.

=========================================================
INTERNAL LOGS (for tracking only)
=========================================================
VERSION: 9.2.2
DATE: auto at runtime
CHANGES:
- v9.2.2  Issue 1 (RTF fallback for resume/cover); keeps v9.2.1 PDF guide + opt-in update checker
- v9.2.1  Issue 2 (real PDF guide if ReportLab), Issue 3 (opt-in update checker)
- v9.2    Mobile-first intro, acronym expander, update-check print, DOCX polish, manifest/logs
- v9.1    Sponsors-ready README & simplified iPhone/Android instructions
- v9.0    Batch scaffolding, risk banner, smarter naming, job-scan JSON

ROADMAP (internal)
- iOS/Android thin clients (SwiftUI/Compose) calling this as API
- Batch playlist mode; background processing
- Font-embedded PDFs; advanced templates
=========================================================
"""

from __future__ import annotations

import os, re, sys, json, difflib, datetime, shutil, subprocess
from pathlib import Path
from typing import List, Dict, Any, Tuple

# ---------------- Optional dependencies ----------------
DOCX_OK = True
try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    DOCX_OK = False

PDF_READ_OK = True
try:
    import PyPDF2  # reading PDFs if available
except Exception:
    PDF_READ_OK = False

REQ_OK = True
try:
    import requests
except Exception:
    REQ_OK = False

# Optional PDF export (pure-Python) — Issue 2
REPORTLAB_OK = True
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.units import inch
except Exception:
    REPORTLAB_OK = False

# ---------------- Versioning / Links ----------------
CURRENT_VERSION = "9.2.2"
RELEASES_PAGE = "https://github.com/hoguej77/hogueresmaster/releases"
LATEST_API = "https://api.github.com/repos/hoguej77/hogueresmaster/releases/latest"
DIRECT_DL = "https://github.com/hoguej77/hogueresmaster/releases/download/v9.2.2/HogueResMaster_v9_2_2.py"

TODAY = datetime.date.today().strftime("%Y-%m-%d")

# ---------------- UX: mobile-first intro ----------------
SHOW_INTRO = True  # set False to suppress console intro

MOBILE_HELP = (
    "\n"
    "📱 Quick Start (ChatGPT / Gemini / Copilot)\n"
    "  1) Upload: resume + job post + this file (.py)\n"
    "  2) Type exactly: Run HogueResMaster\n"
    "  3) You’ll get: Resume (DOCX or RTF), Cover (DOCX or RTF), Interview Guide (HTML + PDF*), ATS text, Redline, Manifest JSON\n"
    "     * RTF when python-docx is missing. PDF guide requires ReportLab; otherwise HTML is provided.\n"
)

def print_intro():
    if not SHOW_INTRO:
        return
    print(f"\\n===== HogueResMaster v{CURRENT_VERSION} =====")
    print("Next‑gen resume & career accelerator — mobile‑ready.")
    print(MOBILE_HELP)
    print("📥 Direct download (current build):", DIRECT_DL)
    print("ℹ️  All releases:", RELEASES_PAGE, "\\n")

# ---------------- Update checker (Issue 3: opt-in) ----------------
def _safe_json(resp):
    try:
        return resp.json()
    except Exception:
        return {}

def check_for_update():
    if not REQ_OK:
        print("ℹ️ Skipping update check (requests not available).")
        return
    try:
        r = requests.get(LATEST_API, timeout=8)
        if not r.ok:
            print("ℹ️ Update check failed (non-200).")
            return
        data = _safe_json(r)
        tag = (data.get("tag_name") or "").strip()
        latest_ver = tag.lstrip("vV")
        def vt(s):
            out = []
            for p in s.split("."):
                try: out.append(int(p))
                except: pass
            return tuple(out)
        if latest_ver and vt(latest_ver) > vt(CURRENT_VERSION):
            print("🔔 Update available!")
            print(f"• Your version: {CURRENT_VERSION}")
            print(f"• Latest:       {latest_ver}  (tag {tag})")
            assets = data.get("assets") or []
            py = [a.get("browser_download_url") for a in assets if str(a.get("browser_download_url","")).endswith(".py")]
            if py:
                print("⬇️  Download latest:", py[0], "\\n")
            else:
                print("🔗 Get it from Releases:", RELEASES_PAGE, "\\n")
        else:
            print("✅ You are on the latest version (or no newer tag found).")
    except Exception as e:
        print(f"⚠️ Could not check for updates: {e}")

# ---------------- Acronym expander ----------------
ACRONYM_GLOSSARY = {
    "ATS": "Applicant Tracking System",
    "STAR": "Situation, Task, Action, Result",
    "CV": "Curriculum Vitae",
    "EMR": "Electronic Medical Record",
    "EHR": "Electronic Health Record",
    "HIPAA": "Health Insurance Portability and Accountability Act",
    "OSHA": "Occupational Safety and Health Administration",
    "JCAHO": "Joint Commission on Accreditation of Healthcare Organizations",
    "PDF": "Portable Document Format",
    "HTML": "HyperText Markup Language",
    "URL": "Uniform Resource Locator",
    "API": "Application Programming Interface",
    "JSON": "JavaScript Object Notation",
    "KPI": "Key Performance Indicator",
}

def expand_acronyms_once(text: str, seen: set | None = None, enabled: bool = True) -> str:
    if not enabled or not text:
        return text
    seen = seen if seen is not None else set()
    acr = "|".join(sorted(map(re.escape, ACRONYM_GLOSSARY.keys()), key=len, reverse=True))
    if not acr:
        return text
    pattern = re.compile(rf"\\b({acr})\\b")
    def repl(m):
        tok = m.group(1)
        if tok in seen:
            return tok
        seen.add(tok)
        return f"{ACRONYM_GLOSSARY.get(tok, tok)} ({tok})"
    return pattern.sub(repl, text)

# ---------------- Scoring / ATS heuristics ----------------
ACTION_VERBS = ["led","owned","built","delivered","implemented","orchestrated","spearheaded",
                "optimized","designed","launched","scaled","improved","reduced","increased",
                "managed","developed","drove","partnered","enabled","accelerated","transformed","modernized"]

STOPWORDS = set("a an and or the but with for to from on in at as by of be is are was were will would shall should can could into within without among across per plus via than then that this those it its your you we our they them their he she his her who whom which what when where why how".split())

def clean_tokens(text: str) -> List[str]:
    return [t for t in re.findall(r"[A-Za-z]{2,}", (text or "").lower()) if t not in STOPWORDS]

def kw_set(text: str, topn=400) -> set:
    toks = clean_tokens(text)
    seen, out = set(), []
    for t in toks:
        if t not in seen:
            seen.add(t); out.append(t)
    return set(out[:topn])

def ats_score(tx: str) -> int:
    vcount = sum(1 for v in ACTION_VERBS if re.search(rf"\\b{re.escape(v)}\\b", (tx or "").lower()))
    nums = len(re.findall(r"[%$€£]\\s?\\d|\\d{1,3}(?:,\\d{3})*(?:\\.\\d+)?%?", tx or ""))
    urls = len(re.findall(r"https?://\\S+", tx or ""))
    raw = 0.5*min(1.0, vcount/18) + 0.3*min(1.0, nums/12) + 0.2*min(1.0, urls/3)
    return int(round(raw*100))

def composite_scores(jd: str, txt: str):
    jks, tks = kw_set(jd), kw_set(txt)
    cov = int(round(100 * len(jks & tks) / max(1, len(jks))))
    ats = ats_score(txt)
    hire = int(round(0.5*cov + 0.5*ats))
    grade = "A+" if hire>=97 else "A" if hire>=93 else "A-" if hire>=90 else "B+" if hire>=87 else \
            "B" if hire>=83 else "B-" if hire>=80 else "C+" if hire>=77 else "C" if hire>=73 else \
            "C-" if hire>=70 else "D" if hire>=60 else "F"
    return cov, ats, hire, grade

# ---------------- IO helpers ----------------
def read_text_any(path_or_url: str) -> str:
    if not path_or_url: return ""
    p = str(path_or_url)
    try:
        if re.match(r"^https?://", p) and REQ_OK:
            r = requests.get(p, timeout=20)
            if r.ok: return r.text
        path = Path(p)
        if not path.exists(): return ""
        if p.lower().endswith(".txt"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if p.lower().endswith(".docx") and DOCX_OK:
            d = docx.Document(str(path))
            return "\\n".join([pg.text for pg in d.paragraphs])
        if p.lower().endswith(".pdf") and PDF_READ_OK:
            txt = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for pg in reader.pages:
                    txt.append(pg.extract_text() or "")
            return "\\n".join(txt)
    except Exception as e:
        print(f"⚠️ Read failure for {p}: {e}")
    return ""

def payload_to_text(payload)->str:
    lines=[payload.get("summary",""), "Core: "+", ".join(payload.get("core_items",[]))]
    for j in payload.get("experience", []):
        lines.append(f"{j.get('title','')} @ {j.get('company','')} — {j.get('dates','')}")
        for b in j.get("bullets",[]): lines.append("• "+b)
    if payload.get("licenses_certs"): lines.append("Licenses/Certs: " + ", ".join(payload["licenses_certs"]))
    if payload.get("education"): lines.append("Education: " + payload.get("education",""))
    return "\\n".join(lines)

# ---------------- Builders / Optimizers ----------------
VET_CROSSWALK = {
    "11B":"Infantryman → Security/Operations Specialist",
    "88M":"Motor Transport → Logistics/Driver Ops",
    "68W":"Combat Medic → EMT/Patient Care Tech",
    "35F":"Intel Analyst → Business/Threat Intelligence",
    "3D1":"Client Systems → IT Support/Helpdesk",
    "3D2":"Cyber Systems → Systems Admin/SRE",
}

def build_initial_payload(name: str, resume_text: str) -> Dict[str, Any]:
    payload = {
        "name": name or "Candidate",
        "summary": "Professional, impact-focused contributor with a track record of delivering measurable results.",
        "core_items": ["Stakeholder Management","Customer Success","Process Improvement","Data-Driven Decisions"],
        "experience": [{
            "company":"Most Recent Company","title":"Most Recent Title","dates":"[CONFIRM]–Present",
            "bullets":[
                "Owned customer outcomes across a portfolio; delivered Quarterly Business Reviews (QBRs) and renewal plans.",
                "Partnered cross-functionally to drive adoption and reduce churn."
            ]}],
        "education":"Education upon request"
    }
    for code, trans in VET_CROSSWALK.items():
        if re.search(rf"\\b{re.escape(code)}\\b", resume_text):
            payload["core_items"].append(trans)
    return payload

def guess_name_from_text(resume_text: str) -> str:
    if not resume_text: return ""
    lines = [ln.strip() for ln in resume_text.splitlines() if ln.strip()]
    for ln in lines[:5]:
        m = re.match(r"^([A-Z][a-z]+(?:\\s[A-Z]\\.)?(?:\\s[A-Z][a-z]+))$", ln)
        if m: return m.group(1)
    m = re.search(r"Name[:\\-]\\s*([A-Z][a-z]+(?:\\s[A-Z]\\.)?(?:\\s[A-Z][a-z]+))", resume_text)
    return m.group(1) if m else ""

def payload_from_resume(name_hint: str, resume_text: str) -> Dict[str, Any]:
    auto = guess_name_from_text(resume_text)
    final_name = (name_hint or auto or "Candidate").strip()
    return build_initial_payload(final_name, resume_text)

def inject_keywords(payload, jd_text, per_section_caps=(2,2,2)):
    missing = list(kw_set(jd_text) - kw_set(payload_to_text(payload)))
    noise = {"and","with","from","will","work","role","team","customer","customers","clients","provide","years","experience","requirements","must","have"}
    missing = [m for m in missing if m not in noise][:6]
    if not missing: return payload
    sum_cap, core_cap, bullet_cap = per_section_caps
    if sum_cap and missing:
        extra = ", ".join(missing[:sum_cap])
        payload["summary"] = (payload.get("summary","") + (" " if payload.get("summary") else "") + f"Focus: {extra}.").strip()
        missing = missing[sum_cap:]
    if core_cap and missing:
        core = payload.get("core_items", [])
        for t in list(missing)[:core_cap]:
            if t.title() not in core: core.append(t.title())
        payload["core_items"] = core
        missing = missing[core_cap:]
    if bullet_cap and missing and payload.get("experience"):
        b = payload["experience"][0].get("bullets", [])
        add = list(missing)[:bullet_cap]
        b.insert(0, f"Advanced {', '.join(add)} initiatives aligned to role requirements.")
        payload["experience"][0]["bullets"] = b
    return payload

def inject_metric_shell(payload):
    if not payload.get("experience"): return payload
    b = payload["experience"][0].get("bullets", [])
    if not any(re.search(r"\\d", x) for x in b):
        b.insert(0, "Improved adoption by [CONFIRM]% and renewal rate by [CONFIRM]% across assigned accounts.")
        payload["experience"][0]["bullets"] = b
    return payload

def optimize_to_target(payload, jd_text, target=95, max_iters=5):
    for _ in range(max_iters):
        txt = payload_to_text(payload)
        cov, ats, hire, _ = composite_scores(jd_text, txt)
        if hire >= target: break
        payload = inject_keywords(payload, jd_text)
        payload = inject_metric_shell(payload)
    return payload

# ---------------- Medical enhancements ----------------
MED_CERTS = {"CNA","LPN","LVN","RN","BSN","MSN","NP","PA-C","MD","DO","EMT","EMT-B","EMT-I","EMT-P",
             "Paramedic","RRT","CRT","PT","DPT","OT","COTA","SLP","PharmD","RPh","CPhT","CRCST","ARRT",
             "CNMT","CPR","BLS","ACLS","PALS","NRP","TNCC","ENPC","ATLS","RHIT","RHIA","CPC","CCS","CCA","COC","CIC","CHDA","CRCR"}
MED_COMPLIANCE = {"HIPAA","OSHA","JCAHO","The Joint Commission","CLIA","CAP","IRB","GCP","ICH-GCP","FDA","EMA","21 CFR Part 11"}
MED_EMR = {"Epic","EpicCare","Cerner","Oracle Health","Meditech","Allscripts","athenahealth","eClinicalWorks","NextGen","Practice Fusion"}
MED_SPECIALTIES = {"Emergency","ICU","CCU","Telemetry","Oncology","Pediatrics","OB/GYN","Labor & Delivery","Surgery",
                   "Perioperative","PACU","Cath Lab","Radiology","Cardiology","Pulmonology","Nephrology","Psychiatry",
                   "Behavioral Health","Primary Care","Urgent Care","Family Medicine","Internal Medicine","Clinical Research",
                   "Pharmacovigilance","Regulatory Affairs","Revenue Cycle","Medical Billing","Coding"}

def looks_medical(text: str) -> bool:
    t = (text or "").lower()
    needles = ["patient","clinic","hospital","emr","ehr","epic","cerner","meditech","provider","nurse","physician","pharmac","radiology","respiratory","triage","icd-10","cpt","hipaa","clinical trial","inpatient","outpatient","charting","billing","coding","revenue cycle"]
    return any(n in t for n in needles)

def extract_med_certs(text: str):
    found = []
    for c in MED_CERTS:
        if re.search(rf"\\b{re.escape(c)}\\b", text):
            found.append(c)
    seen=set(); out=[]
    for c in found:
        if c not in seen: seen.add(c); out.append(c)
    return out

def infer_med_level(text: str) -> str:
    now = datetime.date.today().year
    yrs = 0
    for m in re.findall(r"\\b(20\\d{2}|19\\d{2})\\b", text):
        y=int(m); 
        if 1970<=y<=now: yrs=max(yrs, now-y)
    for m in re.findall(r"\\b(\\d{1,2})\\s+years?\\b", text.lower()):
        yrs=max(yrs, int(m))
    senior = bool(re.search(r"\\b(lead|charge nurse|supervisor|manager|director|chief|attending)\\b", text, re.I))
    research = bool(re.search(r"\\b(PI|co-?investigator|IRB|clinical trial|protocol|GCP)\\b", text, re.I))
    if yrs>=8 or senior or research: return "advanced"
    if yrs>=3: return "mid"
    return "beginner"

def apply_med_enhancements(payload, resume_text, jd_text, domain_mode="auto"):
    active = (domain_mode=="medical") or (domain_mode=="auto" and looks_medical(resume_text + " " + jd_text))
    if not active: return payload
    level = infer_med_level(resume_text + " " + jd_text)

    payload.setdefault("summary","")
    payload.setdefault("core_items",[])
    payload.setdefault("experience",[])

    starter = {
        "beginner":"Compassionate, detail‑oriented healthcare professional focused on safe, efficient patient care and team collaboration.",
        "mid":"Results‑driven clinician with proven impact in patient outcomes, Electronic Medical Record (EMR) workflows, and cross‑disciplinary care coordination.",
        "advanced":"Senior healthcare leader driving clinical quality, regulatory compliance, and operational excellence across multi‑unit settings."
    }[level]

    if "Healthcare" not in payload["summary"]:
        payload["summary"]= (payload["summary"] + (" " if payload["summary"] else "") + starter + " (Healthcare).").strip()

    both = (resume_text + " " + jd_text)
    certs_found = extract_med_certs(both)
    emr_found = [e for e in MED_EMR if re.search(rf"\\b{re.escape(e)}\\b", both)]
    comp_found = [c for c in MED_COMPLIANCE if re.search(rf"\\b{re.escape(c)}\\b", both)]
    specs_found = [s for s in MED_SPECIALTIES if re.search(rf"\\b{re.escape(s)}\\b", both, re.I)]

    def add_core(items):
        cur=[x for x in payload.get("core_items",[]) if x]
        for x in items:
            if x and x not in cur: cur.append(x)
        payload["core_items"]=cur[:18]

    add_core(certs_found[:6]); add_core(emr_found[:4]); add_core(comp_found[:4]); add_core(specs_found[:6])

    if payload.get("experience"):
        b=payload["experience"][0].get("bullets",[])
        if not any(re.search(r"\\d",x) for x in b):
            shell = {"beginner":"Supported avg. [CONFIRM] patients/day with [CONFIRM]% chart accuracy; assisted with Electronic Medical Record (EMR) documentation (Epic).",
                     "mid":"Reduced patient throughput time by [CONFIRM]% via triage workflow improvements; maintained ≥[CONFIRM]% patient satisfaction.",
                     "advanced":"Led [CONFIRM]-bed unit; improved quality indicators (falls, central line–associated bloodstream infection) by [CONFIRM]%; sustained average length of stay and HCAHPS targets."
                     }[level]
            b.insert(0, shell); payload["experience"][0]["bullets"]=b

    payload["medical_level"]=level; payload["domain_applied"]="medical"
    return payload

# ---------------- Writers ----------------
def _rtf_escape(t: str) -> str:
    return (t or "").replace("\\\\", "\\\\\\\\").replace("{", "\\{").replace("}", "\\}")

def write_rtf_resume(out_path: Path, payload: Dict[str, Any]) -> Path:
    name = _rtf_escape(payload.get("name","Candidate"))
    summary = _rtf_escape(payload.get("summary",""))
    core = ", ".join(payload.get("core_items",[]) or [])
    exp = payload.get("experience",[]) or []
    lic = ", ".join(payload.get("licenses_certs",[]) or [])
    edu = _rtf_escape(payload.get("education","") or "")

    r = []
    r.append(r"{\rtf1\ansi\deff0")
    r.append(r"\fs40\b "+name+r"\b0\par")
    r.append(r"\qc Email | Phone | City, ST | linkedin.com/in/username\par\ql")
    r.append(r"\par\b Summary\b0\par "+summary+r"\par")
    if core:
        r.append(r"\par\b Core Skills\b0\par "+_rtf_escape(core)+r"\par")
    if exp:
        r.append(r"\par\b Experience\b0\par")
        for j in exp:
            head = f"{j.get('title','')} — {j.get('company','')} ({j.get('dates','')})"
            r.append(r"\par\b "+_rtf_escape(head)+r"\b0\par")
            for b in j.get("bullets",[]) or []:
                r.append(r"\pard\li480\u8226? "+_rtf_escape(b)+r"\par\pard\li0")
    if lic:
        r.append(r"\par\b Licenses & Certifications\b0\par "+_rtf_escape(lic)+r"\par")
    if edu:
        r.append(r"\par\b Education\b0\par "+edu+r"\par")
    r.append("}")
    out_path.write_text("\\n".join(r), encoding="utf-8")
    return out_path

def write_rtf_cover(out_path: Path, name:str, company:str, title:str) -> Path:
    body = f"""Dear Hiring Team,

I’m excited to apply for the {title or 'role'} at {company or 'your company'}. My background includes leading customer outcomes, driving adoption, and improving renewal performance. I partner cross-functionally, communicate clearly with executives, and turn data into action.

I would welcome a conversation to discuss how I can deliver value quickly.

Sincerely,
{name}
"""
    r = []
    r.append(r"{\rtf1\ansi\deff0")
    for ln in body.splitlines():
        r.append(_rtf_escape(ln)+r"\par")
    r.append("}")
    out_path.write_text("\\n".join(r), encoding="utf-8")
    return out_path

def write_docx_resume(out_path: Path, payload):
    if not DOCX_OK:
        return write_rtf_resume(out_path.with_suffix(".rtf"), payload)
    doc=docx.Document()
    style=doc.styles['Normal']; style.font.name="Calibri"; style.font.size=Pt(11)

    # Name
    p=doc.add_paragraph(payload.get("name","")); p.runs[0].bold=True; p.runs[0].font.size=Pt(16)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # Contact line (ATS-safe, edit later)
    c=doc.add_paragraph("Email | Phone | City, ST | linkedin.com/in/username")
    c.alignment=WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    doc.add_paragraph().add_run("Summary").bold=True
    doc.add_paragraph(payload.get("summary",""))

    # Core Skills
    core = payload.get("core_items",[])
    if core:
        doc.add_paragraph().add_run("Core Skills").bold=True
        doc.add_paragraph(", ".join(core))

    # Experience
    if payload.get("experience"):
        doc.add_paragraph().add_run("Experience").bold=True
        for j in payload["experience"]:
            run=doc.add_paragraph().add_run(f"{j.get('title','')} — {j.get('company','')} ({j.get('dates','')})"); run.bold=True
            for b in j.get("bullets",[]):
                doc.add_paragraph("• "+b)

    # Licenses/Certs
    if payload.get("licenses_certs"):
        doc.add_paragraph().add_run("Licenses & Certifications").bold=True
        doc.add_paragraph(", ".join(payload["licenses_certs"]))

    # Education
    if payload.get("education"):
        doc.add_paragraph().add_run("Education").bold=True
        doc.add_paragraph(payload.get("education",""))

    doc.save(str(out_path))
    return out_path

def write_docx_cover(out_path: Path, name:str, company:str, title:str):
    content = f"""Dear Hiring Team,

I’m excited to apply for the {title or 'role'} at {company or 'your company'}. My background includes leading customer outcomes, driving adoption, and improving renewal performance. I partner cross-functionally, communicate clearly with executives, and turn data into action.

I would welcome a conversation to discuss how I can deliver value quickly.

Sincerely,
{name}
"""
    if not DOCX_OK:
        return write_rtf_cover(out_path.with_suffix(".rtf"), name, company, title)
    doc=docx.Document(); style=doc.styles['Normal']; style.font.name="Calibri"; style.font.size=Pt(11)
    for ln in content.split("\n"): doc.add_paragraph(ln)
    doc.save(str(out_path)); return out_path

def write_html_guide(out_path: Path, company:str, title:str, scores:dict, jd_snip:str):
    html=f"""<!doctype html><html><meta charset='utf-8'><title>Interview Guide</title>
<style>body{{font-family:system-ui,Segoe UI,Arial;margin:24px;}} .k{{padding:8px 12px;border-radius:8px;background:#f5f5f7;display:inline-block;margin:4px 8px 12px 0}}</style>
<h1>Interview Study Guide</h1>
<p><b>Role:</b> {title or 'N/A'} &nbsp; | &nbsp; <b>Company:</b> {company or 'N/A'}</p>
<div class='k'>🎓 Grade: {scores['grade']}</div>
<div class='k'>🔑 Match: {scores['cov']}%</div>
<div class='k'>⚡ Scan (Applicant Tracking System health): {scores['ats']}%</div>
<div class='k'>💼 Hire Likelihood: {scores['hire']}%</div>
<h2>What to Know</h2>
<ul>
<li>Target the <b>Top 5</b> keywords from the job post in your answers.</li>
<li>Bring <b>two metrics</b> (renewal %, adoption %, annual recurring revenue impact).</li>
<li>Prepare a <b>60-second</b> story using Situation, Task, Action, Result.</li>
</ul>
<h2>Job Snippet</h2>
<pre style='white-space:pre-wrap'>{jd_snip[:1500]}</pre>
</html>"""
    out_path.write_text(html, encoding="utf-8"); return out_path

# Issue 2: Real PDF guide when ReportLab is available
def write_pdf_guide(out_pdf: Path, company: str, title: str, scores: dict, jd_snip: str):
    if not REPORTLAB_OK:
        return None
    doc = SimpleDocTemplate(str(out_pdf), pagesize=letter,
                            leftMargin=0.8*inch, rightMargin=0.8*inch,
                            topMargin=0.9*inch, bottomMargin=0.9*inch)
    styles = getSampleStyleSheet()
    H1 = ParagraphStyle('H1', parent=styles['Heading1'], alignment=TA_LEFT, fontSize=16, spaceAfter=12)
    story = []
    story.append(Paragraph("Interview Study Guide", H1))
    story.append(Paragraph(f"<b>Role:</b> {title or 'N/A'} &nbsp;&nbsp; <b>Company:</b> {company or 'N/A'}", styles['Normal']))
    story.append(Paragraph(f"🎓 Grade: {scores.get('grade','N/A')} &nbsp; 🔑 Match: {scores.get('cov','0')}% &nbsp; ⚡ ATS health: {scores.get('ats','0')}% &nbsp; 💼 Hire: {scores.get('hire','0')}%", styles['Normal']))
    story.append(Spacer(1, 10))
    story.append(Paragraph("<b>What to Know</b>", styles['Heading2']))
    story.append(Paragraph("• Target the <b>Top 5</b> keywords from the job post.", styles['Normal']))
    story.append(Paragraph("• Bring <b>two metrics</b> (renewal %, adoption %, annual recurring revenue impact).", styles['Normal']))
    story.append(Paragraph("• Prepare a <b>60-second</b> story using Situation, Task, Action, Result.", styles['Normal']))
    story.append(Spacer(1, 10))
    story.append(Paragraph("<b>Job Snippet</b>", styles['Heading2']))
    story.append(Paragraph((jd_snip or "")[:1500].replace("\n", "<br/>"), styles['Normal']))
    doc.build(story)
    return out_pdf

def write_redline(out_path: Path, before:str, after:str):
    diff = "\n".join(difflib.unified_diff(before.splitlines(), after.splitlines(), fromfile="before", tofile="after", lineterm=""))
    out_path.write_text(diff, encoding="utf-8"); return out_path

def write_ats_txt(out_path: Path, txt:str):
    pre = ("Applicant Tracking System (ATS) notice: This plain text version is optimized for automated scanners.\n"
           "It avoids tables, text boxes, and complex formatting so your resume can be parsed reliably.\n\n")
    out_path.write_text(pre + expand_acronyms_once(txt, set()), encoding="utf-8"); return out_path

# ---------------- Detect files / Guess JD info ----------------
RESUME_EXT = (".docx",".pdf",".txt",".rtf")
JD_EXT = (".docx",".pdf",".txt",".html",".htm")

def detect_files(prefer_dir: Path) -> Dict[str, Path]:
    candidates = []
    for root in {prefer_dir, Path.cwd()}:
        if not root.exists(): continue
        for p in root.iterdir():
            if p.is_file() and p.suffix.lower() in (RESUME_EXT + JD_EXT):
                if p.name.lower().endswith(".py"): continue
                candidates.append(p)
    resumes = [p for p in candidates if any(k in p.name.lower() for k in ["resume","cv"]) or p.suffix.lower() in (".docx",".pdf")]
    jds = [p for p in candidates if any(k in p.name.lower() for k in ["job","jd","posting","description"]) or p.suffix.lower() in (".txt",".pdf",".html",".htm")]
    pick_resume = max(resumes, key=lambda x: x.stat().st_mtime, default=None)
    pick_jd = max(jds, key=lambda x: x.stat().st_mtime, default=None)
    return {"resume": pick_resume, "jd": pick_jd}

def guess_company_and_title(jd_text: str, jd_path: str="") -> Tuple[str,str]:
    company=""; title=""
    m = re.search(r"Company[:\-]\s*([^\n]+)", jd_text, re.I); 
    if m: company = m.group(1).strip()
    m = re.search(r"(Job\s*Title|Title)[:\-]\s*([^\n]+)", jd_text, re.I)
    if m: title = m.group(2).strip()
    if not title:
        lines=[ln.strip() for ln in jd_text.splitlines() if ln.strip()]
        if lines: title = lines[0][:80]
    if not company and jd_path and re.match(r"^https?://", str(jd_path)):
        host = re.sub(r"^https?://","", str(jd_path)).split("/")[0]
        host = host.split(":")[0]
        parts = host.split(".")
        if len(parts)>=2: company = parts[-2].capitalize()
    for bad in ["Careers","Jobs","Job", "Hiring","Openings"]:
        if title.endswith(bad): title = title.replace(bad,"").strip(" -|")
    return company, title

# ---------------- Main pipeline ----------------
def auto_mode_run(outdir: Path, check_updates: bool):
    outdir.mkdir(parents=True, exist_ok=True)
    found = detect_files(Path("/mnt/data"))
    resume_p, jd_p = found.get("resume"), found.get("jd")
    if not resume_p or not jd_p:
        raise SystemExit("Auto‑mode needs a resume and a job posting in this chat. Please upload both, then re‑run.")

    res_text = read_text_any(str(resume_p))
    jd_text = read_text_any(str(jd_p))

    # Build payload and infer name/company/title
    payload = payload_from_resume("", res_text)
    company, title = guess_company_and_title(jd_text, str(jd_p))

    # Domain enhancements
    payload = apply_med_enhancements(payload, res_text, jd_text, domain_mode="auto")

    # Optimize
    before_txt = payload_to_text(payload)
    payload = optimize_to_target(payload, jd_text, target=95, max_iters=5)
    after_txt = payload_to_text(payload)

    cov, ats, hire, grade = composite_scores(jd_text, after_txt)
    scores = {"cov":cov,"ats":ats,"hire":hire,"grade":grade}

    base = f"{(payload['name']).replace(' ','_')}_{(company or 'Company').replace(' ','_')}_{(title or 'Role').replace(' ','_')}_{TODAY}"

    resume_docx = outdir / f"{base}_RESUME.docx"
    cover_docx  = outdir / f"{base}_COVER.docx"
    guide_html  = outdir / f"{base}_GUIDE.html"
    guide_pdf   = outdir / f"{base}_GUIDE.pdf"
    redline_path= outdir / f"{base}_REDLINE.diff.txt"
    ats_path    = outdir / f"{base}_ATS.txt"

    # Write resume/cover; capture actual output paths (docx or rtf)
    resume_out = write_docx_resume(resume_docx, payload)
    cover_out  = write_docx_cover(cover_docx, payload['name'], company or "", title or "")
    write_html_guide(guide_html, company or "", title or "", scores, jd_text[:2000])
    pdf_made = write_pdf_guide(guide_pdf, company or "", title or "", scores, jd_text[:2000])
    write_redline(redline_path, before_txt, after_txt)
    write_ats_txt(ats_path, after_txt)

    # Manifest & run log
    manifest = {
        "summary": f"{payload['name']} targeting {title or 'Role'} @ {company or 'Company'}",
        "scores": scores,
        "files": {
            "resume": str(resume_out),
            "cover": str(cover_out),
            "guide_html": str(guide_html),
            "guide_pdf": str(guide_pdf) if pdf_made else None,
            "ats_txt": str(ats_path),
            "redline_txt": str(redline_path)
        }
    }
    (outdir / f"{base}_MANIFEST.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    logline = {
        "ts": datetime.datetime.now().isoformat(timespec="seconds"),
        "version": CURRENT_VERSION,
        "name": payload['name'],
        "company": company,
        "title": title,
        "scores": scores
    }
    (outdir / "RUN_LOG.ndjson").open("a", encoding="utf-8").write(json.dumps(logline)+"\n")

    print("\\n===== HogueResMaster v9.2.2 (Auto Mode) =====")
    print(f"📄 Name: {payload['name']}")
    print(f"🏢 Company: {company or '[from JD]'}   🧰 Title: {title or '[from JD]'}")
    print(f"📊 Match: {cov}%   ⚡ ATS Health: {ats}%   💼 Hire: {hire}%   🎓 Grade: {grade}")
    if check_updates:
        check_for_update()
    print("\\nFiles saved:")
    for k, v in manifest["files"].items():
        if v: print(" •", k, "→", v)

def build_argparser():
    import argparse
    ap = argparse.ArgumentParser(prog="HogueResMaster v9.2.2", description="Resume generator & optimizer (mobile‑first)")
    ap.add_argument("--name", default="", help="(Optional) Candidate name. If omitted, auto‑extracted from resume.")
    ap.add_argument("--company", default="", help="(Optional) Target company. If omitted, inferred from Job Description (JD).")
    ap.add_argument("--title", default="", help="(Optional) Target role title. If omitted, inferred from Job Description (JD).")
    ap.add_argument("--job", default="", help="(Optional) JD path/URL. If omitted, auto‑detects from uploads.")
    ap.add_argument("--resume", default="", help="(Optional) Resume path. If omitted, auto‑detects from uploads.")
    ap.add_argument("--outdir", default="./out", help="Output folder")
    # Issue 3: opt-in update checker
    ap.add_argument("--check_updates", default="no", choices=["yes","no"], help="Check GitHub for newer release (default: no)")
    return ap

def main():
    args = build_argparser().parse_args()
    outdir = Path(args.outdir or "./out"); outdir.mkdir(parents=True, exist_ok=True)

    # Auto-mode: no explicit paths provided
    if not args.job and not args.resume:
        auto_mode_run(outdir, check_updates=(args.check_updates=="yes"))
        return

    # Manual mode
    jd_text = read_text_any(args.job)
    res_text = read_text_any(args.resume)
    if not jd_text or not res_text:
        raise SystemExit("Provide --job and --resume OR leave both blank to auto‑detect from uploads.")

    payload = payload_from_resume(args.name, res_text)
    company, title = (args.company or ""), (args.title or "")
    if not company or not title:
        c2,t2 = guess_company_and_title(jd_text, args.job)
        company = company or c2; title = title or t2

    payload = apply_med_enhancements(payload, res_text, jd_text, domain_mode="auto")
    before_txt = payload_to_text(payload)
    payload = optimize_to_target(payload, jd_text, target=95, max_iters=5)
    after_txt = payload_to_text(payload)
    cov, ats, hire, grade = composite_scores(jd_text, after_txt)

    base = f"{(payload['name']).replace(' ','_')}_{(company or 'Company').replace(' ','_')}_{(title or 'Role').replace(' ','_')}_{TODAY}"
    resume_docx = outdir / f"{base}_RESUME.docx"
    cover_docx  = outdir / f"{base}_COVER.docx"
    guide_html  = outdir / f"{base}_GUIDE.html"
    guide_pdf   = outdir / f"{base}_GUIDE.pdf"
    redline_path= outdir / f"{base}_REDLINE.diff.txt"
    ats_path    = outdir / f"{base}_ATS.txt"

    # Capture actual output paths
    resume_out = write_docx_resume(resume_docx, payload)
    cover_out  = write_docx_cover(cover_docx, payload['name'], company or "", title or "")
    write_html_guide(guide_html, company or "", title or "", {"cov":cov,"ats":ats,"hire":hire,"grade":grade}, jd_text[:2000])
    pdf_made = write_pdf_guide(guide_pdf, company or "", title or "", {"cov":cov,"ats":ats,"hire":hire,"grade":grade}, jd_text[:2000])
    write_redline(redline_path, before_txt, after_txt)
    write_ats_txt(ats_path, after_txt)

    print("\\n===== HogueResMaster v9.2.2 =====")
    print(f"🧾 Name: {payload['name']} | 🏢 Company: {company or '[from JD]'} | 🧰 Title: {title or '[from JD]'}")
    print(f"🏷️ Grade: {grade}   🧩 Match: {cov}%   (ATS: {ats}%)   💼 Hire: {hire}%")
    if args.check_updates=="yes":
        check_for_update()
    for p in [resume_out, cover_out, guide_html, (guide_pdf if pdf_made else None), redline_path, ats_path]:
        if p: print(" •", p)

if __name__ == "__main__":
    print_intro()
    # Note: update check is opt-in via --check_updates yes
    main()
